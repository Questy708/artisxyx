#!/usr/bin/env python3
"""
xCelero Labs — 10-Year Workplan (0→1)
Generates a comprehensive, downloadable Word document.
"""
import sys, os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Colors ──
ORANGE = RGBColor(0xFF, 0x4D, 0x00)
DARK = RGBColor(0x0A, 0x0A, 0x0A)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xAA, 0xAA, 0xAA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Document defaults ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.3
style.paragraph_format.space_after = Pt(6)

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Helper functions ──
def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(text, level=1, color=DARK, size=None):
    """Add a heading with custom styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
    return h

def add_para(text, bold=False, italic=False, color=None, size=None, align=None, indent=None, space_after=None):
    """Add a paragraph with styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p

def add_number(text):
    """Add a numbered list item."""
    return doc.add_paragraph(text, style='List Number')

def add_page_break():
    doc.add_page_break()

def add_table(headers, rows, col_widths=None, header_color="FF4D00"):
    """Add a formatted table with colored header."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
    # Data rows
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

def add_divider():
    """Add a horizontal divider."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("◆ ◆ ◆")
    run.font.color.rgb = ORANGE
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("xCELERO LABS")
run.font.size = Pt(42)
run.font.color.rgb = ORANGE
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Critical Technology for Emerging Markets")
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("THE TEN-YEAR WORKPLAN")
run.font.size = Pt(28)
run.font.color.rgb = DARK
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("From Zero to One — 2026 to 2035")
run.font.size = Pt(16)
run.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A detailed operational blueprint for building the venture studio,\ninfrastructure platform, and commercialization engine that will\naccelerate human civilization across twenty domains.")
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CLASSIFICATION: INTERNAL — STRATEGIC\nVERSION: 1.0\nDATE: JUNE 2026")
run.font.size = Pt(10)
run.font.color.rgb = LIGHT_GRAY

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("Table of Contents", level=1, color=ORANGE)
doc.add_paragraph()

toc_items = [
    ("PART I: THE THESIS", ""),
    ("  1. Executive Summary", ""),
    ("  2. The Civilization Thesis", ""),
    ("  3. The Four-Pillar Engine", ""),
    ("  4. The Twenty Domains", ""),
    ("  5. Mission, Vision & Operating Principles", ""),
    ("", ""),
    ("PART II: THE TEN-YEAR ARC", ""),
    ("  6. Year 1 (2026): Foundation — Zero to One", ""),
    ("  7. Year 2 (2027): First Light", ""),
    ("  8. Year 3 (2028): Replication", ""),
    ("  9. Year 4 (2029): Regional Scale", ""),
    ("  10. Year 5 (2030): Continental Dominance", ""),
    ("  11. Year 6 (2031): The Network Effect", ""),
    ("  12. Year 7 (2032): Infrastructure Complete", ""),
    ("  13. Year 8 (2033): The Flywheel Compounds", ""),
    ("  14. Year 9 (2034): Civilization Platform", ""),
    ("  15. Year 10 (2035): The Century Engine", ""),
    ("", ""),
    ("PART III: YEAR ONE — THE DETAILED 0→1", ""),
    ("  16. Year 1: Quarter-by-Quarter", ""),
    ("  17. Q1 2026: Month-by-Month", ""),
    ("  18. Month 1 (January 2026): Week-by-Week", ""),
    ("  19. Week 1 (January 1–7, 2026): Day-by-Day", ""),
    ("", ""),
    ("PART IV: FUNCTIONAL WORKSTREAMS", ""),
    ("  20. Infrastructure — The XEmbassy Network", ""),
    ("  21. Ventures — Building Companies", ""),
    ("  22. Capital — The Funding Stack", ""),
    ("  23. Community — The XCitizen Network", ""),
    ("  24. Programs — The Four Pipelines", ""),
    ("  25. Technology — The Platform", ""),
    ("  26. Operations & Talent", ""),
    ("  27. Legal, Governance & Compliance", ""),
    ("", ""),
    ("PART V: METRICS, BUDGETS & RISKS", ""),
    ("  28. KPI Dashboard — Year by Year", ""),
    ("  29. Financial Projections", ""),
    ("  30. The Milestone Map", ""),
    ("  31. Risk Register", ""),
    ("", ""),
    ("PART VI: APPENDICES", ""),
    ("  Appendix A: Organization Chart", ""),
    ("  Appendix B: Hub Rollout Schedule (190 Hubs)", ""),
    ("  Appendix C: Fund Structure & Capital Stack", ""),
    ("  Appendix D: The Twenty Domains — Deep Dive", ""),
    ("  Appendix E: Glossary of Terms", ""),
]

for item, _ in toc_items:
    if item.startswith("PART"):
        p = add_para(item, bold=True, color=ORANGE, size=12, space_after=4)
    elif item == "":
        doc.add_paragraph()
    else:
        add_para(item, size=10, space_after=2)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PART I: THE THESIS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("PART I: THE THESIS", level=1, color=ORANGE, size=24)
doc.add_paragraph()

# ── Chapter 1: Executive Summary ──
add_heading_styled("1. Executive Summary", level=1, color=DARK)

add_para("xCelero Labs is a venture studio and infrastructure platform building critical technology across emerging markets. This document is the ten-year operational workplan to take xCelero from zero to one — from a concept and a website to a civilization-scale commercialization engine operating across 190 hubs, 39 countries, and 20 domains of human progress.", size=11)

add_para("The thesis is simple and blunt: the next century of human civilization will be built — or lost — in the markets that need its breakthroughs most. The bottlenecks are known (food, water, energy, health, education, finance, transport, and twelve more). The capital exists. The talent exists. What does not exist is the commercialization engine that turns capital and talent into compounding infrastructure at the speed and scale the century demands. xCelero Labs is that engine.", size=11)

add_para("This plan is organized around four pillars — Infrastructure, Ventures, Capital, and Community — fused into a single flywheel. Each pillar feeds the next: infrastructure lowers the cost of building ventures; ventures create the returns that attract capital; capital funds the next wave of infrastructure; community compounds the network effect across all three. The flywheel turns slowly at first, then faster, then vertiginously.", size=11)

add_heading_styled("The Ten-Year Targets (Projected Horizon)", level=2, color=ORANGE)
add_table(
    ["Metric", "Year 1 (2026)", "Year 5 (2030)", "Year 10 (2035)"],
    [
        ["Hubs (XEmbassies)", "1", "40", "190"],
        ["Countries", "1", "15", "39"],
        ["Ventures Built", "3", "250", "5,000+"],
        ["Unicorns (Projected)", "0", "5", "200"],
        ["Capital Under Management", "$2M", "$250M", "$1Bn+"],
        ["XCitizens (Community)", "100", "5,000", "100,000"],
        ["Jobs Created", "15", "2,000", "20,000+"],
        ["Domains Active", "3", "12", "20"],
        ["Program Cohorts/Year", "2", "12", "48"],
        ["TownSquare Users", "500", "50,000", "1,000,000+"],
    ],
    col_widths=[5, 3, 3, 3]
)

add_para("These are not today's numbers. They are the targets xCelero Labs is engineered to hit within a decade. Every section of this plan traces a line from the first day of Year 1 to the last day of Year 10, showing exactly how each target is reached.", italic=True, color=GRAY, size=10)

add_heading_styled("How to Read This Plan", level=2, color=ORANGE)
add_para("This document is approximately 100 pages and is structured in six parts:")
add_bullet("Part I (The Thesis) establishes the why — the civilization-scale problem, the four-pillar solution, and the twenty domains.")
add_bullet("Part II (The Ten-Year Arc) walks year by year from 2026 to 2035, giving each year a theme, a set of objectives, key initiatives, and a budget.")
add_bullet("Part III (Year One Detailed) zooms into 2026 — quarter by quarter, then month by month, then week by week, then day by day for the first week. This is the actual 0→1 playbook.")
add_bullet("Part IV (Functional Workstreams) details each of the eight operational pillars — who runs them, what they do, how they are measured.")
add_bullet("Part V (Metrics, Budgets & Risks) provides the KPI dashboard, the financial model, the milestone map, and the risk register.")
add_bullet("Part VI (Appendices) contains the org chart, the 190-hub rollout schedule, the fund structure, the domain deep dives, and a glossary.")

add_para("This is a living document. It will be revised quarterly. The Year 1 detail is prescriptive; the Year 10 detail is directional. The further out we go, the more the plan becomes a compass rather than a map.", italic=True, color=GRAY)

add_page_break()

# ── Chapter 2: The Civilization Thesis ──
add_heading_styled("2. The Civilization Thesis", level=1, color=DARK)

add_heading_styled("The Stakes", level=2, color=ORANGE)
add_para("By 2050, the Earth will be home to as many as ten billion people. To feed them, we must produce more food in the next four decades than all farmers in history have harvested over the past eight thousand years. One point three billion people have no electricity. Two billion drink water that makes them sick. Six hundred and seventeen million children cannot read. One point four billion adults are locked out of banking. The planet is warming, the clinics are empty, the talent is sealed off from capital by the accident of where it was born.")

add_para("These are not separate problems. They are the twenty bottlenecks on which a civilization of ten billion depends. Each one is a domain. Each domain has a known solution set. What is missing is not the science, the capital, or the talent — it is the commercialization engine that turns all three into deployed, operating, compounding infrastructure at the speed the century demands.")

add_heading_styled("The Gap", level=2, color=ORANGE)
add_para("Capital, infrastructure, and talent sit on one side; the markets that need them sit on the other. The gap between them is not a gap of knowledge or money. It is a gap of commercialization — the rail that turns a lab breakthrough into a deployed product, a deployed product into a company, a company into an industry, and an industry into compounding civilization infrastructure. That rail does not exist in emerging markets. xCelero Labs is building it.")

add_heading_styled("The Engine", level=2, color=ORANGE)
add_para("The engine has four parts, each feeding the next:")
add_bullet("Infrastructure (XEmbassies): physical nodes — campuses, labs, foundries, studios — placed in the markets that need them. Each node lowers the cost of building a venture by providing space, power, tools, and a community of operators on day one.")
add_bullet("Ventures: companies built inside the infrastructure, solving the twenty domain bottlenecks. Each venture is a civilization-scale solution wrapped in a business model.")
add_bullet("Capital: the funding stack — from $500 solidarity tickets to $250K+ anchor positions — that finances the ventures and the infrastructure. Six vehicles, one thesis.")
add_bullet("Community (XCitizens): the network of operators, founders, investors, and mentors that compounds across every hub. The community is the connective tissue that turns isolated efforts into collective momentum.")

add_heading_styled("The Compounding", level=2, color=ORANGE)
add_para("Every new venture strengthens the network. Every new hub compounds the reach. Every new XCitizen adds a node to the nervous system. The flywheel turns: infrastructure enables ventures, ventures generate returns, returns attract capital, capital builds more infrastructure, infrastructure attracts more community, community spawns more ventures. The first turn is slow — years of grinding before the first venture ships. The tenth turn is fast. The hundredth turn is vertiginous. By Year 10, the engine is not a startup; it is a civilization platform.")

add_heading_styled("The Outcome", level=2, color=ORANGE)
add_para("Not a forecast, but a blueprint. The bottlenecks break open, domain by domain, until the future is not something we wait for — but something we ship. Five thousand companies. Two hundred unicorns. One hundred thousand XCitizens. Twenty thousand jobs. Twenty domains. One hundred ninety hubs. Thirty-nine countries. One engine. One century.")

add_page_break()

# ── Chapter 3: The Four-Pillar Engine ──
add_heading_styled("3. The Four-Pillar Engine", level=1, color=DARK)

add_para("The xCelero engine is not a fund, not an accelerator, not a coworking space, and not a community. It is all four, fused. Each pillar is a business in its own right; together, they form a flywheel that none could form alone. This chapter details each pillar — what it is, what it does, how it is measured, and how it feeds the next.")

# Pillar 1
add_heading_styled("Pillar 01 — Infrastructure: The XEmbassy Network", level=2, color=ORANGE)
add_para("What it is: A network of physical nodes — XEmbassies — placed in the markets that need critical technology most. Each XEmbassy is a campus containing some combination of: a foundry (manufacturing), a lab (R&D), a studio (software/design), a commons (coworking/community), and an outpost (last-mile deployment).")
add_para("What it does: It lowers the cost of building a venture from 'impossible' to 'inevitable.' A founder who enters an XEmbassy on day one has power, space, tools, a community of operators, and a route to market — all the things that take a Silicon Valley startup three years and three million dollars to assemble. At xCelero, they assemble on day one.")
add_para("How it is measured: Hub count (target: 190 by Year 10), utilization rate (target: 80%+), ventures spawned per hub per year (target: 5+), cost-per-venture-launched (target: 60% below standalone).")
add_para("How it feeds the next: Infrastructure is the soil. Ventures are the trees. Without the soil, the trees take a decade to grow. With it, they grow in a season.")

# Pillar 2
add_heading_styled("Pillar 02 — Ventures: The Companies", level=2, color=ORANGE)
add_para("What it is: Companies built inside the XEmbassy network, each solving a specific bottleneck in one of the twenty domains. xCelero does not 'invest in' companies — it builds them. The studio model: xCelero provides the idea, the team, the capital, the infrastructure, and the go-to-market. The founder provides the obsession.")
add_para("What it does: It turns a domain bottleneck (e.g., '600M Africans lack electricity') into a venture (e.g., Helios — AI-managed modular solar microgrids) and turns that venture into a deployed, revenue-generating, scaling company. Each venture is a civilization-scale solution wrapped in a business model.")
add_para("How it is measured: Ventures launched (target: 5,000+ by Year 10), survival rate at 3 years (target: 75%), unicorns (target: 200), aggregate revenue (target: $10Bn+ by Year 10), aggregate jobs created (target: 20,000+).")
add_para("How it feeds the next: Ventures generate returns. Returns attract capital. Capital funds the next wave. Without ventures, there is nothing to invest in; without returns, there is no reason to invest.")

# Pillar 3
add_heading_styled("Pillar 03 — Capital: The Funding Stack", level=2, color=ORANGE)
add_para("What it is: Six vehicles, one thesis — back the technology the next century needs, in the markets that need it most. The stack runs from $500 solidarity tickets (democratized access) to $250K+ anchor positions (institutional scale), with everything in between: thematic funds, a venture studio fund, development finance, a non-dilutive grant desk, and a secondary liquidity vehicle.")
add_para("What it does: It matches the capital to the reality of building in the Global South — where a venture might need grant funding for a pilot, equity for scale, debt for working capital, and concessionary capital for market-entry, all in the same year. Solidarity pricing (below-market carry for early LPs) lets the market form before it is squeezed.")
add_para("How it is measured: Capital under management (target: $1Bn+ by Year 10), LP count (target: 5,000+), IRR (target: 25%+ net), DPI by Year 7 (target: 1.5x+), grant capital deployed (target: $200M+).")
add_para("How it feeds the next: Capital builds infrastructure. Infrastructure lowers the cost of the next venture. The next venture generates the next return. The flywheel turns.")

# Pillar 4
add_heading_styled("Pillar 04 — Community: The XCitizen Network", level=2, color=ORANGE)
add_para("What it is: The connective tissue — a network of operators, founders, investors, and mentors spanning every hub. The XCitizen program runs cohorts of 100 per quarter; each cohort is a cross-pollinated group of people who run infrastructure, build ventures, deploy capital, and transfer knowledge. TownSquare is the digital forum where the network convenes between cohorts.")
add_para("What it does: It compounds network effects. A founder in Lagos can find a mentor in Nairobi, an investor in Cape Town, and a pilot site in Kano — all within the network. The community is the reason a venture built at xCelero scales faster than one built alone. It is also the reason xCelero itself scales: every XCitizen is a node that attracts more nodes.")
add_para("How it is measured: XCitizen count (target: 100,000+ by Year 10), active engagement rate (target: 60% monthly), cross-hub collaborations per quarter (target: 500+), mentor hours per venture (target: 100+ in Year 1).")
add_para("How it feeds the next: Community spawns ventures. Ventures generate returns. Returns attract capital. Capital builds infrastructure. Infrastructure attracts community. The flywheel is complete.")

add_divider()
add_para("The four pillars are not sequential. They are simultaneous. They start on Day 1 of Year 1 and compound for ten years. The rest of this plan is the operational detail of how.", italic=True, color=GRAY)

add_page_break()

# ── Chapter 4: The Twenty Domains ──
add_heading_styled("4. The Twenty Domains", level=1, color=DARK)

add_para("xCelero Labs builds ventures across twenty domains — the pillars of a civilization. Each domain has a bottleneck (the 'Now') and a solution set (the 'Future'). The studio does not operate in all twenty on Day 1; it sequences them, starting with the three where the bottleneck is most acute and the solution is most deployable. The full twenty are active by Year 10.")

add_table(
    ["#", "Domain", "The Bottleneck (Now)", "The Solution (Future)", "Active From"],
    [
        ["1", "Food", "More food needed in 40 years than in 8,000", "Vertical farms, AI crops, lab protein", "Year 1"],
        ["2", "Energy", "1.3B without power; diesel skies", "Fusion, solar paint, peer-to-peer grids", "Year 1"],
        ["3", "Water", "2B drink sick water; child dies every 80s", "Atmospheric generators, low-energy desal", "Year 1"],
        ["4", "Health", "1 in 5 have diagnostics; find out too late", "Hospital of the future: sensor + AI", "Year 2"],
        ["5", "Longevity", "73 yrs avg; $1T/yr managing decline", "Cellular reprogramming, senolytics", "Year 4"],
        ["6", "Education", "617M illiterate; syllabus a century behind", "1:1 AI tutor per child, free, $20 phone", "Year 2"],
        ["7", "Finance", "1.4B unbanked; 15% remittance fees", "Money as internet protocol; instant", "Year 1"],
        ["8", "Transport", "75% of product cost is logistics", "eVTOL, hyperloop, one multimodal OS", "Year 3"],
        ["9", "Cities", "70% urban by 2050; slums fastest-growing", "Cities that grow like coral; self-powering", "Year 3"],
        ["10", "The Moon", "12 walked it; none since 1972", "Permanent settlement; industry off-world", "Year 7"],
        ["11", "Mars", "0 humans; 30M-mile desert", "Self-sustaining city of 1M; backup world", "Year 8"],
        ["12", "Humanoid Robots", "$100T/yr of manual human effort", "$20K robot per home; work optional", "Year 5"],
        ["13", "Intelligence", "Frontier model costs $100M; concentrated", "Compute as utility; AI on every phone", "Year 2"],
        ["14", "Manufacturing", "Sell lithium, buy back battery at 10×", "Micro-factories; 12-block supply chain", "Year 2"],
        ["15", "Materials", "Mine-to-landfill one-way pipe", "Atom-by-atom design; mine the landfill", "Year 4"],
        ["16", "Connectivity", "2.6B offline; invisible to digital economy", "Satellite blanket; last billion online", "Year 2"],
        ["17", "Governance", "17th-century state for a satellite planet", "Governance as protocol; opt-in communities", "Year 6"],
        ["18", "Defense", "$2.4T/yr buying hardware for last war", "Defense as software; war irrational", "Year 5"],
        ["19", "Science", "0.1% of GDP on basic research", "AI runs 1B experiments in simulation", "Year 3"],
        ["20", "Knowledge", "Talent sealed off by birthplace", "Single network; talent finds capital in a day", "Year 1"],
    ],
    col_widths=[1, 3, 4.5, 4.5, 2.5]
)

add_para("The sequencing is deliberate. Year 1 starts with Food, Energy, Water, Finance, and Knowledge — the domains where the bottleneck is most acute, the solution is most deployable, and the venture can reach revenue fastest. Health, Education, Intelligence, Manufacturing, and Connectivity follow in Year 2. The frontier domains (Moon, Mars, Humanoid Robots) come later, when the engine has the capital and the talent to attempt them.", size=10)

add_page_break()

# ── Chapter 5: Mission, Vision, Operating Principles ──
add_heading_styled("5. Mission, Vision & Operating Principles", level=1, color=DARK)

add_heading_styled("Mission", level=2, color=ORANGE)
add_para("To build the commercialization engine that accelerates human civilization across twenty domains — turning capital and talent into deployed, compounding infrastructure at the speed the century demands.")

add_heading_styled("Vision", level=2, color=ORANGE)
add_para("By 2035, xCelero Labs operates 190 hubs across 39 countries, has built 5,000+ ventures (200 projected unicorns), manages $1Bn+ in capital, and connects 100,000+ XCitizens in a single nervous system for a century. The bottlenecks break open, domain by domain, until the future is not something we wait for — but something we ship.")

add_heading_styled("Operating Principles", level=2, color=ORANGE)

principles = [
    ("Build, don't bet.", "xCelero is a studio, not a fund. We build companies; we do not merely invest in them. The difference is ownership, control, and speed. We own the idea, the team, the capital, and the infrastructure from day one."),
    ("Infrastructure first.", "Before we build a venture, we build the soil it grows in. An XEmbassy lowers the cost of the next ten ventures by 60%. The infrastructure is the moat."),
    ("Solidarity pricing.", "The markets we build in cannot bear Silicon Valley economics. Our capital is priced to let the market form before it is squeezed. Early LPs get below-market carry; ventures get patient capital. The returns come from scale, not extraction."),
    ("Talent over credentials.", "A genius in a village without a road is a genius who never gets to be one. We find talent where it lives, not where it went to school. The XCitizen network is the largest talent scout on earth."),
    ("Compounding over velocity.", "We are not building a startup; we are building a century. Short-term speed matters less than long-term compounding. Every decision is weighed against: does this make the flywheel turn faster in Year 10?"),
    ("Radical transparency.", "The plan, the numbers, the failures, and the returns are visible to every XCitizen. Trust is the currency of networks; opacity is the tax on it."),
    ("Domain obsession.", "We do not chase trends. We chase bottlenecks. Every venture must map to one of the twenty domains. If it does not accelerate civilization in a measurable way, we do not build it."),
    ("The route is the strategy.", "Geography is not an afterthought; it is the thesis. We build along routes — the Gulf of Guinea Arc, the Sahel, the East African Corridor — because civilization infrastructure compounds along trade routes, not in isolation."),
]
for title, body in principles:
    add_para(title, bold=True, color=DARK, size=12)
    add_para(body, size=11, indent=0.5)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PART II: THE TEN-YEAR ARC
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("PART II: THE TEN-YEAR ARC", level=1, color=ORANGE, size=24)
doc.add_paragraph()
add_para("This part walks year by year from 2026 to 2035. Each year has a theme, a set of strategic objectives, key initiatives by quarter, a budget, and the KPIs that define success. The detail is prescriptive for Years 1-3 and directional for Years 4-10 — the further out we go, the more the plan becomes a compass.", italic=True, color=GRAY)

add_page_break()

# ── Helper for year chapters ──
def add_year_chapter(year_num, year_label, theme, tagline, objectives, q_initiatives, budget_lines, kpis, narrative):
    add_heading_styled(f"{year_num}. Year {year_label}: {theme}", level=1, color=DARK)
    add_para(tagline, italic=True, color=ORANGE, size=14)
    doc.add_paragraph()

    add_heading_styled("Strategic Objectives", level=2, color=ORANGE)
    for obj in objectives:
        add_bullet(obj)

    add_heading_styled("Key Initiatives by Quarter", level=2, color=ORANGE)
    for q, initiatives in q_initiatives:
        add_para(q, bold=True, color=DARK, size=12)
        for item in initiatives:
            add_bullet(item, level=1)
        doc.add_paragraph()

    add_heading_styled("Budget Allocation", level=2, color=ORANGE)
    add_table(["Workstream", "Allocation (USD)", "% of Total"], budget_lines, col_widths=[6, 4, 3])
    doc.add_paragraph()

    add_heading_styled("KPI Targets — End of Year", level=2, color=ORANGE)
    add_table(["KPI", "Target"], kpis, col_widths=[8, 5])
    doc.add_paragraph()

    add_heading_styled("Narrative", level=2, color=ORANGE)
    add_para(narrative)
    add_page_break()

# ═══ YEAR 1 (2026) ═══
add_year_chapter(
    6, "1 (2026)", "Foundation — Zero to One",
    "The year we stop talking and start building. One hub. Three ventures. One hundred XCitizens. Two program cohorts. The engine's first turn.",
    [
        "Establish the legal entity, the founding team, and the operating systems.",
        "Open XEmbassy Nairobi (Hub 01) — the first physical node.",
        "Launch 3 venture studios (Energy, Water, Finance domains).",
        "Run 2 program cohorts (xHansa Fellowship Q1, xCelero Accelerator Q3).",
        "Close the Studio Fund I at $2M (seed capital for ventures + operations).",
        "Onboard the first 100 XCitizens and launch TownSquare.",
        "Build the technology platform (website, forum, internal CRM).",
    ],
    [
        ("Q1 2026 (January–March): The Setup Quarter", [
            "Incorporate xCelero Labs (Mauritius HoldCo + Kenya OpCo).",
            "Secure founding team: CEO, CTO, Head of Ventures, Head of Capital, Head of Community, Head of Operations (6 people).",
            "Sign the lease on XEmbassy Nairobi (M1 Core, 1,200 sqm).",
            "Build the website (already prototyped) and launch publicly.",
            "Open the Studio Fund I raise ($2M target).",
            "Recruit Cohort 1 of the xHansa Fellowship (25 fellows).",
            "Begin venture ideation across 5 domains (Food, Energy, Water, Finance, Knowledge).",
        ]),
        ("Q2 2026 (April–June): The Launch Quarter", [
            "XEmbassy Nairobi opens (April 1).",
            "Run xHansa Fellowship Cohort 1 (8-week crucible).",
            "Spin out Venture 01 (Energy — Helios: solar microgrids).",
            "Spin out Venture 02 (Water — Nimbus: atmospheric water generators).",
            "Launch TownSquare (the digital forum) with the first 100 XCitizens.",
            "First Demo Day (June) — 5 xHansa ventures present to investors.",
            "Close Studio Fund I at $2M.",
        ]),
        ("Q3 2026 (July–September): The Validation Quarter", [
            "Launch xCelero Accelerator Cohort 1 (10 ventures, 20 founders).",
            "Spin out Venture 03 (Finance — cross-border payments).",
            "Deploy Helios pilot: 5 microgrid nodes in Kano, Nigeria.",
            "Deploy Nimbus pilot: 2 AWG units in Zinder, Niger.",
            "Open applications for xHansa Fellowship Cohort 2.",
            "Hire: Head of Infrastructure, 2 venture builders, 1 community manager.",
            "Begin site selection for Hub 02 (Lagos) and Hub 03 (Accra).",
        ]),
        ("Q4 2026 (October–December): The Compounding Quarter", [
            "Run xCelero Accelerator Demo Day (December) — 10 ventures to investors.",
            "Helios pilot: reach 50 nodes, first revenue ($5K MRR).",
            "Nimbus pilot: reach 10 units, first municipal contract.",
            "Onboard XCitizens 101–200 (Cohort 2).",
            "Sign lease on XEmbassy Lagos (Hub 02) — opening Q1 2027.",
            "Close the year with: 1 hub, 3 ventures, 200 XCitizens, $2M raised, 15 jobs created.",
            "Board meeting: review Year 1, approve Year 2 plan and budget.",
        ]),
    ],
    [
        ["Team & Operations", "$420,000", "21%"],
        ["Infrastructure (Nairobi hub)", "$500,000", "25%"],
        ["Venture Building (3 studios)", "$600,000", "30%"],
        ["Programs (2 cohorts)", "$200,000", "10%"],
        ["Technology & Platform", "$150,000", "7.5%"],
        ["Community & Events", "$80,000", "4%"],
        ["Legal & Compliance", "$50,000", "2.5%"],
        ["Total", "$2,000,000", "100%"],
    ],
    [
        ["Hubs open", "1 (Nairobi)"],
        ["Ventures launched", "3"],
        ["Ventures with first revenue", "2 (Helios, Nimbus)"],
        ["Program cohorts run", "2 (xHansa + Accelerator)"],
        ["XCitizens onboarded", "200"],
        ["Capital raised (Studio Fund I)", "$2M"],
        ["Jobs created", "15"],
        ["Domains active", "3 (Energy, Water, Finance)"],
        ["TownSquare users", "500"],
        ["Pilot deployments", "62 (50 microgrids + 10 AWG + 2 finance)"],
    ],
    "Year 1 is the hardest year. Nothing compounds yet; everything is being built for the first time. The temptation is to move fast and break things; the discipline is to build the infrastructure — legal, operational, cultural — that the next nine years will stand on. By December 2026, the engine has turned once: one hub, three ventures, two cohorts, two hundred XCitizens. The flywheel exists. It is slow. But it exists."
)

# ═══ YEAR 2 (2027) ═══
add_year_chapter(
    7, "2 (2027)", "First Light",
    "The first ventures reach revenue. The second hub opens. The engine proves it can replicate. The flywheel turns twice.",
    [
        "Open Hubs 02 (Lagos) and 03 (Accra) — the Gulf of Guinea Arc begins.",
        "Scale Helios to 500 nodes; Nimbus to 100 units; Finance venture to 10K users.",
        "Launch 5 new ventures (Health, Education, Manufacturing, Connectivity, Intelligence).",
        "Run 4 program cohorts (2 xHansa + 2 Accelerator).",
        "Raise Studio Fund II ($10M) and the first Thematic Fund ($5M, Energy).",
        "Grow XCitizen network to 1,000.",
        "Hire the second layer of leadership (VP-level across each pillar).",
    ],
    [
        ("Q1 2027", [
            "Open XEmbassy Lagos (Hub 02).",
            "Helios reaches 200 nodes, $20K MRR.",
            "Launch Venture 04 (Health — distributed diagnostics).",
            "Run xHansa Fellowship Cohort 3 (30 fellows, first multi-hub cohort).",
            "Begin Studio Fund II raise ($10M).",
        ]),
        ("Q2 2027", [
            "Open XEmbassy Accra (Hub 03).",
            "Launch Venture 05 (Education — AI tutor platform) and Venture 06 (Manufacturing — micro-factory).",
            "Run xCelero Accelerator Cohort 2 (15 ventures).",
            "Close the Energy Thematic Fund at $5M.",
            "TownSquare reaches 2,000 users.",
        ]),
        ("Q3 2027", [
            "Launch Venture 07 (Connectivity — last-mile satellite) and Venture 08 (Intelligence — edge compute).",
            "Helios reaches 500 nodes; first cross-border deployment (Tanzania).",
            "Run xHansa Fellowship Cohort 4 (40 fellows, 3 hubs).",
            "Close Studio Fund II at $10M.",
        ]),
        ("Q4 2027", [
            "Run xCelero Accelerator Cohort 3 (20 ventures) and Demo Day.",
            "Nimbus reaches 100 units; first contract with a municipal government.",
            "Finance venture reaches 10K active users; first profitable month.",
            "Year-end: 3 hubs, 8 ventures, 1,000 XCitizens, $17M raised, 60 jobs.",
            "Board meeting: approve Year 3 plan, the Sahel route expansion.",
        ]),
    ],
    [
        ["Team & Operations", "$1,500,000", "15%"],
        ["Infrastructure (3 hubs)", "$2,000,000", "20%"],
        ["Venture Building (8 studios)", "$4,000,000", "40%"],
        ["Programs (4 cohorts)", "$800,000", "8%"],
        ["Technology & Platform", "$700,000", "7%"],
        ["Community & Events", "$500,000", "5%"],
        ["Capital Raising", "$300,000", "3%"],
        ["Legal & Compliance", "$200,000", "2%"],
        ["Total", "$10,000,000", "100%"],
    ],
    [
        ["Hubs open", "3 (Nairobi, Lagos, Accra)"],
        ["Ventures launched", "8 (cumulative)"],
        ["Ventures with revenue", "5"],
        ["Program cohorts (cumulative)", "6"],
        ["XCitizens", "1,000"],
        ["Capital raised (cumulative)", "$17M"],
        ["Jobs created", "60"],
        ["Domains active", "8"],
        ["TownSquare users", "2,000"],
        ["Aggregate venture MRR", "$80K"],
    ],
    "Year 2 is when the engine proves it can replicate. Hub 01 was a bespoke build; Hubs 02 and 03 test whether the XEmbassy model can be templated. It can. The first ventures reach revenue — not much, but real. The first cross-border deployment happens. The first Thematic Fund closes. The flywheel has turned twice, and the second turn was faster than the first. The pattern is established: each year, the engine does more, faster, with less friction."
)

# ═══ YEAR 3 (2028) ═══
add_year_chapter(
    8, "3 (2028)", "Replication",
    "The Gulf of Guinea Arc is complete. The Sahel route begins. The studio proves it can build across ten domains. The flywheel turns three times — and the third turn is twice as fast as the first.",
    [
        "Open Hubs 04–07 (Kano, Abidjan, Addis Ababa, Kampala) — 7 hubs total.",
        "Launch 7 new ventures (Food, Transport, Cities, Science, Materials, Longevity, Governance).",
        "Scale existing ventures: Helios 2,000 nodes, Nimbus 500 units, Finance 50K users.",
        "Run 6 program cohorts across 4 programs.",
        "Raise Studio Fund III ($25M) + 2 Thematic Funds (Health $10M, Agri $10M).",
        "Grow XCitizens to 3,000; TownSquare to 10,000 users.",
        "Launch the Inception Studios program (ideation → ProtoCo → NewCo).",
        "First venture crosses $1M ARR.",
    ],
    [
        ("Q1 2028", [
            "Open Hubs 04 (Kano) and 05 (Abidjan).",
            "Launch Venture 09 (Food — vertical farms) and Venture 10 (Transport — logistics OS).",
            "Helios reaches 1,000 nodes; first $100K MRR month.",
            "Launch Inception Studios Cohort 1 (5 ProtoCos).",
            "Begin the Sahel route: site selection for Niamey, Bamako, Ouagadougou.",
        ]),
        ("Q2 2028", [
            "Open Hubs 06 (Addis Ababa) and 07 (Kampala).",
            "Launch Venture 11 (Cities — modular housing) and Venture 12 (Science — AI research).",
            "Run xCelero Accelerator Cohorts 4 & 5 (30 ventures total).",
            "Close Health Thematic Fund ($10M) and Agri Thematic Fund ($10M).",
            "First venture (Finance) crosses $1M ARR.",
        ]),
        ("Q3 2028", [
            "Launch Venture 13 (Materials — programmable matter) and Venture 14 (Longevity — senolytics).",
            "Run xHansa Fellowship Cohorts 5 & 6 (80 fellows across 7 hubs).",
            "Begin Studio Fund III raise ($25M).",
            "TownSquare reaches 10,000 users; first community-led venture spinout.",
        ]),
        ("Q4 2028", [
            "Launch Venture 15 (Governance — digital identity).",
            "Run Inception Studios Demo Day; 3 ProtoCos become NewCos.",
            "Helios reaches 2,000 nodes; Nimbus 500 units.",
            "Year-end: 7 hubs, 15 ventures, 3,000 XCitizens, $62M raised, 200 jobs.",
            "Board approves Year 4: East African Corridor + Southern Africa.",
        ]),
    ],
    [
        ["Team & Operations", "$3,000,000", "12%"],
        ["Infrastructure (7 hubs)", "$5,000,000", "20%"],
        ["Venture Building (15 studios)", "$10,000,000", "40%"],
        ["Programs (6 cohorts)", "$2,000,000", "8%"],
        ["Technology & Platform", "$1,500,000", "6%"],
        ["Community & Events", "$1,000,000", "4%"],
        ["Capital Raising", "$1,500,000", "6%"],
        ["Legal & Compliance", "$1,000,000", "4%"],
        ["Total", "$25,000,000", "100%"],
    ],
    [
        ["Hubs open", "7"],
        ["Ventures launched (cumulative)", "15"],
        ["Ventures with $1M+ ARR", "1"],
        ["Program cohorts (cumulative)", "12"],
        ["XCitizens", "3,000"],
        ["Capital raised (cumulative)", "$62M"],
        ["Jobs created", "200"],
        ["Domains active", "10"],
        ["TownSquare users", "10,000"],
        ["Aggregate venture MRR", "$500K"],
    ],
    "Year 3 is when the studio model is proven at scale. Fifteen ventures across ten domains. Seven hubs across two routes. The first $1M ARR venture. The first community-led spinout. The Inception Studios program launches — the ideation pipeline that will feed the accelerator for the next seven years. The engine is no longer an experiment; it is a machine. The question shifts from 'can we build one venture?' to 'how many can we build in parallel?' The answer, by Year 4, is: a lot."
)

add_page_break()

# ═══ YEARS 4-10 (condensed but detailed) ═══
add_year_chapter(
    9, "4 (2029)", "Regional Scale",
    "The East African Corridor opens. The studio hits 30 ventures. The first unicorn is born.",
    [
        "Open Hubs 08–14 (Dar es Salaam, Kigali, Lusaka, Harare, Maputo, Cape Town, Johannesburg) — 14 hubs total.",
        "Launch 10 new ventures; total reaches 25.",
        "Run 10 program cohorts across all 4 programs.",
        "Raise Studio Fund IV ($50M) + 3 Thematic Funds ($30M total).",
        "Grow XCitizens to 6,000; TownSquare to 30,000 users.",
        "First venture crosses $10M ARR; first unicorn valuation ($1Bn).",
        "Open the Quest Fellowship (university partnership program, Cohort 1).",
    ],
    [
        ("Q1 2029", ["Open Hubs 08–10 (Dar, Kigali, Lusaka).", "Launch 3 ventures (Robotics-adjacent, Defense-adjacent, Space-adjacent R&D).", "Helios reaches 5,000 nodes.", "Quest Fellowship Cohort 1 (20 students, Queen's University + African partners)."]),
        ("Q2 2029", ["Open Hubs 11–14 (Harare, Maputo, Cape Town, Jo'burg).", "Launch 4 ventures.", "Run 3 program cohorts.", "First venture crosses $10M ARR (Finance)."]),
        ("Q3 2029", ["Launch 3 ventures.", "First unicorn: Finance venture valued at $1Bn in Series B.", "Close Studio Fund IV ($50M).", "TownSquare reaches 30,000 users."]),
        ("Q4 2029", ["Year-end review: 14 hubs, 25 ventures, 6,000 XCitizens, $142M raised, 500 jobs.", "Board approves Year 5: Southern Africa + first North Africa hub.", "Begin planning the Moon/Mars R&D vertical (long-horizon)."]),
    ],
    [
        ["Team & Operations", "$6,000,000", "12%"],
        ["Infrastructure (14 hubs)", "$10,000,000", "20%"],
        ["Venture Building (25 studios)", "$20,000,000", "40%"],
        ["Programs (10 cohorts)", "$4,000,000", "8%"],
        ["Technology & Platform", "$3,000,000", "6%"],
        ["Community & Events", "$2,000,000", "4%"],
        ["Capital Raising", "$3,000,000", "6%"],
        ["Legal & Compliance", "$2,000,000", "4%"],
        ["Total", "$50,000,000", "100%"],
    ],
    [
        ["Hubs open", "14"],
        ["Ventures (cumulative)", "25"],
        ["Unicorns", "1"],
        ["XCitizens", "6,000"],
        ["Capital raised (cumulative)", "$142M"],
        ["Jobs created", "500"],
        ["Domains active", "13"],
        ["TownSquare users", "30,000"],
    ],
    "Year 4 is the year of the first unicorn. The Finance venture — cross-border payments for the unbanked — hits a $1Bn valuation in its Series B. It is the proof point that xCelero-built companies can reach category-leading scale. The East African Corridor is complete. The Southern Africa route begins. The studio is now building 10 ventures per year and running 10 cohorts. The engine is a factory."
)

add_year_chapter(
    10, "5 (2030)", "Continental Dominance",
    "Halfway. The engine covers sub-Saharan Africa. The first North Africa hub opens. Twenty hubs. Two hundred fifty ventures. Five unicorns.",
    [
        "Open Hubs 15–20 (Cairo, Tunis, Algiers, Casablanca, Dakar, Douala) — 20 hubs, 4 routes active.",
        "Launch 15 new ventures; total reaches 40.",
        "Run 12 program cohorts.",
        "Raise Studio Fund V ($100M) + 4 Thematic Funds ($80M total).",
        "Grow XCitizens to 10,000; TownSquare to 75,000 users.",
        "5 ventures cross $100M ARR; 5 unicorns total.",
        "Launch the Humanoid Robotics R&D vertical (long-horizon bet).",
        "First xCelero venture IPO (Nairobi Securities Exchange).",
    ],
    [
        ("Q1 2030", ["Open Hubs 15–16 (Cairo, Tunis) — North Africa route begins.", "Launch 4 ventures.", "First venture IPO (NSE listing).", "Helios reaches 20,000 nodes; 4 countries."]),
        ("Q2 2030", ["Open Hubs 17–18 (Algiers, Casablanca).", "Launch 4 ventures.", "Run 4 program cohorts.", "3 more unicorns; total reaches 4."]),
        ("Q3 2030", ["Open Hubs 19–20 (Dakar, Douala).", "Launch 4 ventures.", "Humanoid Robotics vertical launches (3 R&D ProtoCos).", "Close Studio Fund V ($100M)."]),
        ("Q4 2030", ["Launch 3 ventures.", "5th unicorn.", "Year-end: 20 hubs, 40 ventures, 10,000 XCitizens, $322M raised, 2,000 jobs.", "Board approves Year 6: the push to 40 hubs and global expansion."]),
    ],
    [
        ["Team & Operations", "$12,000,000", "12%"],
        ["Infrastructure (20 hubs)", "$20,000,000", "20%"],
        ["Venture Building (40 studios)", "$40,000,000", "40%"],
        ["Programs (12 cohorts)", "$8,000,000", "8%"],
        ["Technology & Platform", "$6,000,000", "6%"],
        ["Community & Events", "$4,000,000", "4%"],
        ["Capital Raising", "$6,000,000", "6%"],
        ["Legal & Compliance", "$4,000,000", "4%"],
        ["Total", "$100,000,000", "100%"],
    ],
    [
        ["Hubs open", "20"],
        ["Ventures (cumulative)", "40"],
        ["Unicorns", "5"],
        ["IPOs", "1"],
        ["XCitizens", "10,000"],
        ["Capital raised (cumulative)", "$322M"],
        ["Jobs created", "2,000"],
        ["Domains active", "15"],
        ["TownSquare users", "75,000"],
    ],
    "Year 5 is the halfway mark and the inflection point. The engine covers all of Africa — 20 hubs across 4 routes. The first IPO happens. Five unicorns. The Humanoid Robotics vertical launches — the first long-horizon bet that will not pay off for five years but that positions xCelero for the second half of the decade. The flywheel is now turning fast enough that the studio can build 15 ventures in a year and have 5 of them reach unicorn scale. The compounding is visible."
)

add_page_break()

# Years 6-10 (directional but detailed)
add_year_chapter(
    11, "6 (2031)", "The Network Effect",
    "The engine goes global. First hubs outside Africa. The XCitizen network crosses 20,000. The flywheel is now self-sustaining.",
    [
        "Open Hubs 21–30: first hubs in the Middle East (Dubai, Riyadh), South Asia (Mumbai, Bangalore), Southeast Asia (Jakarta, Singapore) — 30 hubs.",
        "Launch 20 new ventures; total reaches 60.",
        "Run 16 program cohorts across 6 regions.",
        "Raise Studio Fund VI ($150M) + first Global Thematic Fund ($100M).",
        "Grow XCitizens to 20,000; TownSquare to 150,000 users.",
        "10 unicorns; 2 more IPOs.",
        "First cross-continental venture (operates in Africa + Asia simultaneously).",
    ],
    [
        ("Q1 2031", ["Open Hubs 21–23 (Dubai, Riyadh, Mumbai).", "Launch 5 ventures.", "10th unicorn."]),
        ("Q2 2031", ["Open Hubs 24–26 (Bangalore, Jakarta, Singapore).", "Launch 5 ventures.", "Run 6 cohorts across 3 continents."]),
        ("Q3 2031", ["Open Hubs 27–30 (4 more hubs in Asia/ME).", "Launch 5 ventures.", "Close Global Thematic Fund ($100M)."]),
        ("Q4 2031", ["Launch 5 ventures.", "2 more IPOs.", "Year-end: 30 hubs, 60 ventures, 20,000 XCitizens, $572M raised, 5,000 jobs."]),
    ],
    [
        ["Team & Operations", "$18,000,000", "12%"],
        ["Infrastructure (30 hubs)", "$30,000,000", "20%"],
        ["Venture Building", "$60,000,000", "40%"],
        ["Programs", "$12,000,000", "8%"],
        ["Technology", "$9,000,000", "6%"],
        ["Community", "$6,000,000", "4%"],
        ["Capital Raising", "$9,000,000", "6%"],
        ["Legal & Compliance", "$6,000,000", "4%"],
        ["Total", "$150,000,000", "100%"],
    ],
    [
        ["Hubs open", "30"],
        ["Ventures (cumulative)", "60"],
        ["Unicorns", "10"],
        ["IPOs (cumulative)", "3"],
        ["XCitizens", "20,000"],
        ["Capital raised (cumulative)", "$572M"],
        ["Jobs created", "5,000"],
        ["Domains active", "17"],
        ["TownSquare users", "150,000"],
    ],
    "Year 6 is when xCelero stops being an African story and becomes a global one. The first hubs outside Africa test whether the model travels. It does — the XEmbassy template is domain-agnostic and geography-agnostic. The XCitizen network crosses 20,000, and the network effect becomes self-sustaining: ventures now find their co-founders, investors, and pilot sites through the community, not through xCelero HQ. The engine is running itself."
)

add_year_chapter(
    12, "7 (2032)", "Infrastructure Complete",
    "The African route network is finished — all 190 hub locations are identified, 50 are open. The Moon R&D vertical produces its first patent.",
    [
        "Open Hubs 31–50 (Latin America begins: São Paulo, Mexico City, Bogotá, Lima, Buenos Aires) — 50 hubs.",
        "Launch 25 new ventures; total reaches 85.",
        "Run 20 program cohorts.",
        "Raise Studio Fund VII ($200M).",
        "Grow XCitizens to 35,000; TownSquare to 300,000 users.",
        "15 unicorns; 5 IPOs.",
        "Moon R&D vertical: first patent filed (in-situ resource utilization).",
        "First venture crosses $1Bn ARR.",
    ],
    [
        ("Q1 2032", ["Open Hubs 31–35 (Latin America).", "Launch 6 ventures.", "First $1Bn ARR venture (Finance)."]),
        ("Q2 2032", ["Open Hubs 36–42.", "Launch 6 ventures.", "Moon vertical: first patent.", "Run 8 cohorts."]),
        ("Q3 2032", ["Open Hubs 43–50.", "Launch 7 ventures.", "15th unicorn.", "Close Studio Fund VII ($200M)."]),
        ("Q4 2032", ["Launch 6 ventures.", "2 more IPOs.", "Year-end: 50 hubs, 85 ventures, 35,000 XCitizens, $772M raised, 8,000 jobs."]),
    ],
    [
        ["Team & Operations", "$24,000,000", "12%"],
        ["Infrastructure (50 hubs)", "$40,000,000", "20%"],
        ["Venture Building", "$80,000,000", "40%"],
        ["Programs", "$16,000,000", "8%"],
        ["Technology", "$12,000,000", "6%"],
        ["Community", "$8,000,000", "4%"],
        ["Capital Raising", "$12,000,000", "6%"],
        ["Legal & Compliance", "$8,000,000", "4%"],
        ["Total", "$200,000,000", "100%"],
    ],
    [
        ["Hubs open", "50"],
        ["Ventures (cumulative)", "85"],
        ["Unicorns", "15"],
        ["IPOs (cumulative)", "5"],
        ["XCitizens", "35,000"],
        ["Capital raised (cumulative)", "$772M"],
        ["Jobs created", "8,000"],
        ["Domains active", "18"],
        ["TownSquare users", "300,000"],
        ["Patents filed", "1 (Moon ISRU)"],
    ],
    "Year 7 is when the infrastructure buildout crosses the tipping point. 50 hubs open — a quarter of the eventual 190. The Moon R&D vertical files its first patent, a small step toward off-world industry but a proof point that xCelero can operate on decade-long horizons. The first venture crosses $1Bn ARR. The engine is now producing unicorns at a rate of 5 per year. The compounding is vertiginous."
)

add_page_break()

add_year_chapter(
    13, "8 (2033)", "The Flywheel Compounds",
    "100 hubs. 150 ventures. The Mars R&D vertical launches. The engine is now the largest venture studio in the world.",
    [
        "Open Hubs 51–100 (full coverage of Africa, Asia, Latin America; first European hubs) — 100 hubs.",
        "Launch 40 new ventures; total reaches 125.",
        "Run 28 program cohorts.",
        "Raise Studio Fund VIII ($300M).",
        "Grow XCitizens to 50,000; TownSquare to 500,000 users.",
        "30 unicorns; 10 IPOs.",
        "Mars R&D vertical launches (in partnership with a launch provider).",
        "First humanoid robotics venture reaches production.",
        "Aggregate venture revenue crosses $5Bn.",
    ],
    [
        ("Q1 2033", ["Open Hubs 51–65.", "Launch 10 ventures.", "Mars R&D vertical launches.", "20th unicorn."]),
        ("Q2 2033", ["Open Hubs 66–80.", "Launch 10 ventures.", "First humanoid robotics venture in production.", "Run 12 cohorts."]),
        ("Q3 2033", ["Open Hubs 81–100.", "Launch 10 ventures.", "30th unicorn.", "Close Studio Fund VIII ($300M)."]),
        ("Q4 2033", ["Launch 10 ventures.", "3 more IPOs.", "Year-end: 100 hubs, 125 ventures, 50,000 XCitizens, $1.07Bn raised, 12,000 jobs."]),
    ],
    [
        ["Team & Operations", "$36,000,000", "12%"],
        ["Infrastructure (100 hubs)", "$60,000,000", "20%"],
        ["Venture Building", "$120,000,000", "40%"],
        ["Programs", "$24,000,000", "8%"],
        ["Technology", "$18,000,000", "6%"],
        ["Community", "$12,000,000", "4%"],
        ["Capital Raising", "$18,000,000", "6%"],
        ["Legal & Compliance", "$12,000,000", "4%"],
        ["Total", "$300,000,000", "100%"],
    ],
    [
        ["Hubs open", "100"],
        ["Ventures (cumulative)", "125"],
        ["Unicorns", "30"],
        ["IPOs (cumulative)", "10"],
        ["XCitizens", "50,000"],
        ["Capital raised (cumulative)", "$1.07Bn"],
        ["Jobs created", "12,000"],
        ["Domains active", "19"],
        ["TownSquare users", "500,000"],
        ["Aggregate venture revenue", "$5Bn"],
    ],
    "Year 8 is the year the engine becomes the largest venture studio in the world by every measure — hubs, ventures, capital, community. The Mars vertical launches, a bet that will not pay off for a decade but that positions xCelero as the only studio on earth building across all twenty domains, including off-world. The humanoid robotics venture reaches production — the first $20K humanoid, built in an XEmbassy foundry. The flywheel is now producing more value per year than the first five years combined."
)

add_year_chapter(
    14, "9 (2034)", "Civilization Platform",
    "150 hubs. 250 ventures. The engine is no longer a studio; it is a platform — the operating system for a century.",
    [
        "Open Hubs 101–150 (full global coverage; first North American hubs) — 150 hubs.",
        "Launch 50 new ventures; total reaches 175.",
        "Run 36 program cohorts.",
        "Raise Studio Fund IX ($400M).",
        "Grow XCitizens to 75,000; TownSquare to 750,000 users.",
        "50 unicorns; 20 IPOs.",
        "The Connectivity domain is solved: last billion online via xCelero satellite constellation venture.",
        "First venture crosses $10Bn ARR.",
        "Aggregate venture revenue crosses $20Bn.",
    ],
    [
        ("Q1 2034", ["Open Hubs 101–120.", "Launch 12 ventures.", "Connectivity venture: last billion online.", "40th unicorn."]),
        ("Q2 2034", ["Open Hubs 121–135.", "Launch 12 ventures.", "First $10Bn ARR venture.", "Run 16 cohorts."]),
        ("Q3 2034", ["Open Hubs 136–150.", "Launch 13 ventures.", "50th unicorn.", "Close Studio Fund IX ($400M)."]),
        ("Q4 2034", ["Launch 13 ventures.", "5 more IPOs.", "Year-end: 150 hubs, 175 ventures, 75,000 XCitizens, $1.47Bn raised, 16,000 jobs."]),
    ],
    [
        ["Team & Operations", "$48,000,000", "12%"],
        ["Infrastructure (150 hubs)", "$80,000,000", "20%"],
        ["Venture Building", "$160,000,000", "40%"],
        ["Programs", "$32,000,000", "8%"],
        ["Technology", "$24,000,000", "6%"],
        ["Community", "$16,000,000", "4%"],
        ["Capital Raising", "$24,000,000", "6%"],
        ["Legal & Compliance", "$16,000,000", "4%"],
        ["Total", "$400,000,000", "100%"],
    ],
    [
        ["Hubs open", "150"],
        ["Ventures (cumulative)", "175"],
        ["Unicorns", "50"],
        ["IPOs (cumulative)", "20"],
        ["XCitizens", "75,000"],
        ["Capital raised (cumulative)", "$1.47Bn"],
        ["Jobs created", "16,000"],
        ["Domains active", "20 (all)"],
        ["TownSquare users", "750,000"],
        ["Aggregate venture revenue", "$20Bn"],
    ],
    "Year 9 is when all twenty domains are active. The Connectivity venture — a satellite constellation — brings the last billion people online. The engine is now a platform: other studios, funds, and governments plug into the xCelero network to build ventures, deploy capital, and access infrastructure. The first venture crosses $10Bn ARR. xCelero is no longer a startup; it is the operating system for a century."
)

add_year_chapter(
    15, "10 (2035)", "The Century Engine",
    "190 hubs. 39 countries. 5,000 ventures. 200 unicorns. $1Bn+ capital. 100,000 XCitizens. The engine is complete. The century begins.",
    [
        "Open the final 40 hubs (Hubs 151–190) — full global network.",
        "Launch the final wave of ventures; total reaches 5,000+ (including community-led spinouts).",
        "Run 48 program cohorts.",
        "Close the Decade Fund ($500M) — the permanent capital vehicle.",
        "Grow XCitizens to 100,000; TownSquare to 1,000,000+ users.",
        "200 unicorns; 50 IPOs.",
        "Aggregate venture revenue crosses $100Bn.",
        "The Moon venture deploys its first permanent habitat module.",
        "xCelero transitions from a studio to a trust — the permanent steward of the engine.",
    ],
    [
        ("Q1 2035", ["Open Hubs 151–170.", "100th unicorn.", "Moon venture: first habitat module deployed.", "Close Decade Fund ($500M)."]),
        ("Q2 2035", ["Open Hubs 171–185.", "150th unicorn.", "Run 16 cohorts.", "1,000,000th TownSquare user."]),
        ("Q3 2035", ["Open Hubs 186–190 (final 5).", "200th unicorn.", "Aggregate venture revenue crosses $100Bn.", "xCelero Trust established."]),
        ("Q4 2035", ["Decade celebration + 10-year retrospective.", "Publish the next 10-year plan (2036–2045).", "Year-end: 190 hubs, 5,000+ ventures, 100,000 XCitizens, $1.97Bn raised, 20,000+ jobs."]),
    ],
    [
        ["Team & Operations", "$60,000,000", "12%"],
        ["Infrastructure (190 hubs)", "$100,000,000", "20%"],
        ["Venture Building", "$200,000,000", "40%"],
        ["Programs", "$40,000,000", "8%"],
        ["Technology", "$30,000,000", "6%"],
        ["Community", "$20,000,000", "4%"],
        ["Capital Raising", "$30,000,000", "6%"],
        ["Legal & Compliance", "$20,000,000", "4%"],
        ["Total", "$500,000,000", "100%"],
    ],
    [
        ["Hubs open", "190 (complete)"],
        ["Countries", "39"],
        ["Ventures (cumulative)", "5,000+"],
        ["Unicorns", "200"],
        ["IPOs (cumulative)", "50"],
        ["XCitizens", "100,000"],
        ["Capital raised (cumulative)", "$1.97Bn"],
        ["Jobs created", "20,000+"],
        ["Domains active", "20 (all)"],
        ["TownSquare users", "1,000,000+"],
        ["Aggregate venture revenue", "$100Bn"],
        ["Moon habitat", "1 module deployed"],
    ],
    "Year 10 is the end of the beginning. The engine is complete — 190 hubs, 39 countries, 20 domains, 5,000 ventures, 200 unicorns, 100,000 XCitizens. The flywheel is now permanent: it turns whether or not anyone pushes it. xCelero transitions from a studio to a trust — a permanent steward of the engine, governed by the XCitizen network, capitalized by the Decade Fund, and oriented toward the next century. The Moon venture deploys its first habitat. The Mars venture launches its first crewed mission prep. The next ten-year plan (2036–2045) is published. The century begins."
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PART III: YEAR ONE — THE DETAILED 0→1
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("PART III: YEAR ONE — THE DETAILED 0→1", level=1, color=ORANGE, size=24)
doc.add_paragraph()
add_para("This part zooms into 2026 — the first year, the zero-to-one. It breaks the year into quarters, Q1 into months, Month 1 into weeks, and Week 1 into days. This is the actual playbook: what happens, when, by whom, with what budget. If the rest of this plan is a compass, this part is a map.", italic=True, color=GRAY)
add_page_break()

# ── Chapter 16: Year 1 Quarter-by-Quarter (summary table) ──
add_heading_styled("16. Year 1: Quarter-by-Quarter Summary", level=1, color=DARK)
add_para("Year 1 has four quarters, each with a theme and a set of deliverables. The table below is the at-a-glance summary; the sections that follow break Q1 into months, and Month 1 into weeks and days.")
add_table(
    ["Quarter", "Theme", "Key Deliverables", "Budget"],
    [
        ["Q1 (Jan–Mar)", "The Setup", "Incorporate, hire founding team, sign Nairobi lease, launch website, open Fund I raise, recruit xHansa Cohort 1, begin venture ideation", "$420K"],
        ["Q2 (Apr–Jun)", "The Launch", "XEmbassy Nairobi opens, xHansa Cohort 1 runs, Venture 01 (Helios) + 02 (Nimbus) spin out, TownSquare launches, first Demo Day, Fund I closes ($2M)", "$550K"],
        ["Q3 (Jul–Sep)", "The Validation", "Accelerator Cohort 1 launches, Venture 03 (Finance) spins out, Helios pilot (50 nodes), Nimbus pilot (10 units), hire 4 more, site selection for Lagos/Accra", "$550K"],
        ["Q4 (Oct–Dec)", "The Compounding", "Accelerator Demo Day, Helios first revenue, XCitizen Cohort 2, Lagos lease signed, year-end review, Year 2 plan approved", "$480K"],
    ],
    col_widths=[3, 3, 7, 3]
)
add_page_break()

# ── Chapter 17: Q1 2026 Month-by-Month ──
add_heading_styled("17. Q1 2026: Month-by-Month", level=1, color=DARK)
add_para("Q1 is the setup quarter — three months to go from a concept and a prototype to a legal entity, a founding team, a physical hub, a public presence, a fund raise, a recruited cohort, and a venture pipeline. Everything that follows in Year 1 depends on Q1 being executed cleanly.")

# January
add_heading_styled("January 2026: Incorporate & Assemble", level=2, color=ORANGE)
add_para("Theme: The legal and human foundation. By January 31, xCelero Labs is a legal entity with a founding team, a signed lease, and a public website.", bold=True)
add_para("Deliverables:")
add_bullet("Incorporate xCelero Labs HoldCo in Mauritius (international holding structure).")
add_bullet("Register xCelero Kenya OpCo (operating entity for Hub 01).")
add_bullet("Open bank accounts (Mauritius, Kenya).")
add_bullet("Sign the lease on M1 Core Nairobi (1,200 sqm, 5-year term).")
add_bullet("Hire the founding 6: CEO, CTO, Head of Ventures, Head of Capital, Head of Community, Head of Operations.")
add_bullet("Launch the public website (already prototyped — goes live January 15).")
add_bullet("Draft and circulate the Studio Fund I PPM (Private Placement Memorandum).")
add_bullet("Open applications for xHansa Fellowship Cohort 1 (target: 100 applicants for 25 seats).")
add_bullet("Begin venture ideation sessions (5 domains: Food, Energy, Water, Finance, Knowledge).")
add_para("Budget: $140,000 (legal $30K, payroll $70K, lease deposit $20K, website $10K, misc $10K)")

# February
add_heading_styled("February 2026: Build & Recruit", level=2, color=ORANGE)
add_para("Theme: The operational buildout. By February 28, the hub is being fitted out, the team is operational, the fund raise is live, and the first cohort is being selected.", bold=True)
add_para("Deliverables:")
add_bullet("Begin fit-out of XEmbassy Nairobi (partitions, power, network, furniture).")
add_bullet("Set up internal systems: CRM (Notion/Airtable), accounting (Xero), comms (Slack + email).")
add_bullet("First investor meetings for Studio Fund I (target: 20 meetings, 5 soft commitments).")
add_bullet("Close xHansa Cohort 1 applications; conduct 50 interviews; select 25 fellows.")
add_bullet("Deep-dive venture ideation: select 3 venture concepts for Q2 spinout (Helios, Nimbus, + 1).")
add_bullet("Hire: 2 venture-builders (one per venture concept), 1 designer, 1 engineer for the platform team.")
add_bullet("Draft the xCelero Accelerator curriculum (12-week program).")
add_bullet("Begin community building: first 20 XCitizens invited (founders, operators, mentors from the network).")
add_para("Budget: $150,000 (payroll $90K, fit-out $40K, tools $5K, travel $10K, misc $5K)")

# March
add_heading_styled("March 2026: Prepare & Pre-Launch", level=2, color=ORANGE)
add_para("Theme: The final preparations before the April launch. By March 31, the hub is ready, the cohort is confirmed, the ventures are scoped, and the fund has its first close.", bold=True)
add_para("Deliverables:")
add_bullet("Complete XEmbassy Nairobi fit-out; final inspections; ready for April 1 opening.")
add_bullet("First close of Studio Fund I (target: $500K from 5-10 LPs).")
add_bullet("Onboard xHansa Cohort 1 fellows (25); pre-program orientation.")
add_bullet("Finalize Venture 01 (Helios) and Venture 02 (Nimbus) scope: team, budget, pilot plan, 12-month milestones.")
add_bullet("Hire: Hub manager (Nairobi), 2 more venture-builders, 1 community manager.")
add_bullet("Soft-launch TownSquare (invite-only, 50 users — the first XCitizens).")
add_bullet("Press strategy: line up 3-5 media pieces for the April launch.")
add_bullet("Board formation: identify and invite 3-5 board members (first meeting in April).")
add_para("Budget: $130,000 (payroll $80K, final fit-out $30K, fund-raising travel $10K, events $5K, misc $5K)")

add_page_break()

# ── Chapter 18: Month 1 (January 2026) Week-by-Week ──
add_heading_styled("18. Month 1 (January 2026): Week-by-Week", level=1, color=DARK)
add_para("January is the most important month of the decade. Everything that follows is built on what happens in these 31 days. This section breaks January into four weeks, each with specific daily deliverables.")

# Week 1
add_heading_styled("Week 1 (January 1–7): The First Week", level=2, color=ORANGE)
add_para("Theme: Incorporate, assemble the team, and make the first decisions. This is the week the concept becomes a company.")
week1_days = [
    ("Thursday, January 1", "New Year's Day (holiday). No operations. The founding CEO reviews the 10-year plan one final time and commits."),
    ("Friday, January 2", "File incorporation papers for xCelero Labs HoldCo (Mauritius). Engage Mauritius legal counsel. Begin the Kenya OpCo registration. Open the founding team Slack workspace. First all-hands (virtual): CEO + the 5 committed co-founders. Agenda: assign Week 1 owners for each workstream."),
    ("Saturday, January 3", "CEO and Head of Operations travel to Nairobi (if not already there). In-person site visit of M1 Core (the prospective hub). Meeting with the building landlord. Begin lease negotiation."),
    ("Sunday, January 4", "Head of Capital begins drafting the Studio Fund I PPM. CTO begins the final website review (the prototype is built; this is the polish pass). Head of Community begins the XCitizen outreach list (target: first 100 names)."),
    ("Monday, January 5", "Lease negotiation continues. CEO and Head of Ops meet with a Kenya corporate lawyer to finalize OpCo registration. Head of Ventures begins the venture ideation framework (5 domains × 3 concepts each = 15 concepts to evaluate). CTO deploys the website to production (soft launch — no announcement yet)."),
    ("Tuesday, January 6", "First formal team meeting (in-person, Nairobi). All 6 founders present. Agenda: (1) final org structure, (2) Q1 budget approval, (3) hiring plan, (4) venture ideation kickoff, (5) Fund I timeline. Meeting outputs: each founder has a 90-day OKR sheet."),
    ("Wednesday, January 7", "Lease signed (target). Head of Ops opens Kenya bank account (preliminary, pending OpCo registration). Head of Community sends the first 20 XCitizen invitations. CTO sets up the internal CRM and project management tools. Week 1 review: all founders report progress against OKRs."),
]
for day, activity in week1_days:
    add_para(day, bold=True, color=DARK, size=11)
    add_para(activity, size=10, indent=0.5)

# Week 2
add_heading_styled("Week 2 (January 8–14): Systems & Strategy", level=2, color=ORANGE)
add_para("Theme: Set up the operational systems and begin the strategic work — fund raise, venture ideation, cohort recruitment.")
add_bullet("January 8–9: Incorporation filed (Mauritius + Kenya). Bank accounts opened. Accounting system configured.")
add_bullet("January 10: Studio Fund I PPM first draft circulated to the founding team for review.")
add_bullet("January 11: First investor meeting (warm introduction — angel investor in the network).")
add_bullet("January 12: xHansa Fellowship Cohort 1 applications open publicly (website + social + network).")
add_bullet("January 13: Venture ideation session #1 (Energy domain — 3 concepts generated: Helios microgrids, Ember off-grid kits, Ignis fusion broker).")
add_bullet("January 14: Week 2 review. All systems operational. 5 investor meetings booked. 15 cohort applications received.")

# Week 3
add_heading_styled("Week 3 (January 15–21): Public & Pipeline", level=2, color=ORANGE)
add_para("Theme: The public launch and the pipeline buildout.")
add_bullet("January 15: WEBSITE GOES LIVE. Public announcement (LinkedIn, Twitter, press list). The 10-year plan is published as a downloadable document.")
add_bullet("January 16: First wave of press inquiries. CEO does 3-5 interviews (TechCrunch Africa, local Kenyan press, African tech podcasts).")
add_bullet("January 17: Venture ideation session #2 (Water domain — 3 concepts: Nimbus AWG, desal, purification).")
add_bullet("January 18: Investor meetings (target: 5 this week). First soft commitment for Fund I ($50K).")
add_bullet("January 19: xHansa applications cross 50. Head of Community begins interviewing the first batch.")
add_bullet("January 20: Venture ideation session #3 (Finance domain — 3 concepts: cross-border payments, micro-credit, trade finance).")
add_bullet("January 21: Week 3 review. Website: 5,000 unique visitors. Applications: 60. Investor meetings: 8 total. Soft commitments: $100K.")

# Week 4
add_heading_styled("Week 4 (January 22–31): Solidify & Select", level=2, color=ORANGE)
add_para("Theme: Solidify the foundations. Select the first ventures. Narrow the cohort. Prepare for February.")
add_bullet("January 22–23: Venture ideation sessions #4 and #5 (Food + Knowledge domains). Total concepts generated: 15 across 5 domains.")
add_bullet("January 24: Venture selection meeting. The founding team selects 3 concepts for Q2 spinout: Helios (Energy), Nimbus (Water), and a Finance concept (TBD).")
add_bullet("January 25: First venture-builder hired (for Helios). Offer extended to a second (for Nimbus).")
add_bullet("January 26–28: Cohort interviews continue. 30 interviews conducted. Shortlist of 35 for 25 seats.")
add_bullet("January 29: Fund I PPM finalized and sent to 15 prospective LPs.")
add_bullet("January 30: Hub fit-out begins (construction team on-site at M1 Core).")
add_bullet("January 31: MONTH 1 REVIEW. Metrics: entity incorporated, team of 6 (+ 2 hires), lease signed, website live (10K visitors), 80 cohort applications, $150K soft commitments, 3 venture concepts selected, hub fit-out started. On track for Q1.")

add_page_break()

# ── Chapter 19: Week 1 Day-by-Day (already covered above, add more detail) ──
add_heading_styled("19. Week 1 (January 1–7, 2026): Day-by-Day Detail", level=1, color=DARK)
add_para("This section expands on Week 1 with the specific actions, owners, and decision points for each day. Week 1 is the template for how every week should be run: clear ownership, daily check-ins, and a bias toward action.")

add_table(
    ["Day", "Date", "Owner", "Key Action", "Decision/Output"],
    [
        ["Wed", "Jan 1", "CEO", "Review 10-year plan; final commit", "Go/No-Go decision: GO"],
        ["Thu", "Jan 2", "CEO + All", "File Mauritius incorporation; open Slack; first all-hands", "Entity filed; team aligned on Week 1 OKRs"],
        ["Fri", "Jan 3", "CEO + Head of Ops", "Travel to Nairobi; site visit M1 Core; begin lease negotiation", "Site approved; lease terms drafted"],
        ["Sat", "Jan 4", "Head of Capital", "Draft Studio Fund I PPM (first draft)", "PPM v1 circulated to team"],
        ["Sat", "Jan 4", "CTO", "Final website review; deploy to staging", "Website ready for Jan 15 launch"],
        ["Sat", "Jan 4", "Head of Community", "Build XCitizen outreach list (100 names)", "List of 100; first 20 to invite Week 2"],
        ["Sun", "Jan 5", "CEO + Lawyer", "Kenya OpCo registration; bank account prep", "Registration filed; bank meeting booked"],
        ["Sun", "Jan 5", "Head of Ventures", "Launch venture ideation framework (5 domains)", "Framework ready; 15 concepts to evaluate"],
        ["Mon", "Jan 6", "All (in-person)", "First formal team meeting; 90-day OKRs assigned", "Each founder has OKR sheet; Q1 budget approved"],
        ["Tue", "Jan 7", "All", "Lease signed; bank account opened; Week 1 review", "Lease executed; all systems go for Week 2"],
    ],
    col_widths=[1.5, 2, 3, 5, 5]
)

add_para("By the end of Week 1, xCelero Labs is a legal entity with a signed lease, a deployed website (staging), a founding team with clear OKRs, a Fund I PPM in draft, a venture ideation framework, and a list of 100 prospective XCitizens. The concept is now a company. Week 2 begins the operational buildout.", italic=True, color=GRAY)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PART IV: FUNCTIONAL WORKSTREAMS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("PART IV: FUNCTIONAL WORKSTREAMS", level=1, color=ORANGE, size=24)
doc.add_paragraph()
add_para("This part details each of the eight operational workstreams — the functions that run every day, across every year, and that together constitute the xCelero engine. Each workstream has an owner, a mandate, a set of processes, and a measurement framework.", italic=True, color=GRAY)
add_page_break()

# ── Chapter 20: Infrastructure ──
add_heading_styled("20. Infrastructure — The XEmbassy Network", level=1, color=DARK)
add_heading_styled("Owner: Head of Infrastructure (VP-level from Year 2)", level=2, color=ORANGE)
add_para("Mandate: Build and operate the physical network of XEmbassies — the campuses, labs, foundries, studios, and commons that lower the cost of building ventures by 60%.")

add_heading_styled("The XEmbassy Typology", level=2, color=ORANGE)
add_table(
    ["Type", "Size", "Components", "Cost to Build", "Cost to Run/yr", "Ventures Supported"],
    [
        ["M1 Core (Flagship)", "1,000–1,500 sqm", "Foundry + Lab + Studio + Commons + Outpost", "$500K", "$200K", "10–15"],
        ["M2 Node (Regional)", "500–800 sqm", "Studio + Commons + Outpost", "$250K", "$100K", "5–8"],
        ["M3 Outpost (Last-mile)", "100–200 sqm", "Outpost + mini-Studio", "$80K", "$40K", "2–3"],
    ],
    col_widths=[3, 2.5, 4, 2.5, 2.5, 2.5]
)

add_heading_styled("Hub Selection Criteria", level=2, color=ORANGE)
add_para("Every hub location is evaluated against five criteria:")
add_bullet("Route alignment: Is the location on one of the six routes (Gulf of Guinea, Sahel, East African, Central, Southern, North Africa)?")
add_bullet("Market density: Is there a population of 1M+ within 100km who need the domains xCelero builds in?")
add_bullet("Talent supply: Is there a university or technical workforce within commuting distance?")
add_bullet("Regulatory feasibility: Can xCelero operate (register entities, import equipment, deploy pilots) without undue friction?")
add_bullet("Community readiness: Are there 20+ prospective XCitizens in the area to seed the hub?")

add_heading_styled("Rollout Cadence", level=2, color=ORANGE)
add_table(
    ["Year", "New Hubs", "Cumulative", "Route Focus", "Type Mix"],
    [
        ["1", "1", "1", "East African", "1× M1"],
        ["2", "2", "3", "Gulf of Guinea", "2× M1"],
        ["3", "4", "7", "Gulf of Guinea + East African", "2× M1, 2× M2"],
        ["4", "7", "14", "East African + Southern", "2× M1, 3× M2, 2× M3"],
        ["5", "6", "20", "North Africa + West", "2× M1, 2× M2, 2× M3"],
        ["6", "10", "30", "Middle East + South Asia", "3× M1, 4× M2, 3× M3"],
        ["7", "20", "50", "Latin America + SE Asia", "5× M1, 8× M2, 7× M3"],
        ["8", "50", "100", "Global", "10× M1, 20× M2, 20× M3"],
        ["9", "50", "150", "Global", "10× M1, 20× M2, 20× M3"],
        ["10", "40", "190", "Global (complete)", "8× M1, 16× M2, 16× M3"],
    ],
    col_widths=[1.5, 2, 2, 4, 4]
)

add_heading_styled("Operational Standards", level=2, color=ORANGE)
add_bullet("Every hub opens with: power (renewable where possible), internet (100Mbps+), security (24/7), and a hub manager on-site.")
add_bullet("Every hub maintains 80%+ utilization (measured monthly: desks occupied, lab hours used, foundry hours run).")
add_bullet("Every hub runs a weekly community night (open to all XCitizens in the area) and a monthly demo day (ventures present progress).")
add_bullet("Every hub reports to HQ weekly: utilization, ventures in residence, community engagement, issues.")

add_page_break()

# ── Chapter 21: Ventures ──
add_heading_styled("21. Ventures — Building Companies", level=1, color=DARK)
add_heading_styled("Owner: Head of Ventures (VP-level)", level=2, color=ORANGE)
add_para("Mandate: Build ventures that solve the twenty domain bottlenecks. The studio model: xCelero provides the idea, the team, the capital, the infrastructure, and the go-to-market. The founder provides the obsession.")

add_heading_styled("The Venture Pipeline", level=2, color=ORANGE)
add_para("Ventures flow through a five-stage pipeline, each stage gated by specific criteria:")
add_table(
    ["Stage", "Duration", "What Happens", "Gate to Next Stage", "Capital Deployed"],
    [
        ["1. Ideation", "2–4 weeks", "Domain analysis, concept generation, market sizing, feasibility study", "Concept approved by venture committee", "$0"],
        ["2. ProtoCo", "4–8 weeks", "Build the MVP, validate with 5-10 customers, test the unit economics", "5+ customer interviews, viable unit economics", "$25K"],
        ["3. NewCo", "3–6 months", "Incorporate, hire founding team, deploy pilot, reach first revenue", "Pilot deployed, first revenue, team of 3+", "$100K"],
        ["4. Scale", "6–24 months", "Scale the team, expand geographically, raise external round", "$500K+ ARR or equivalent milestone", "$500K"],
        ["5. Spinout", "Ongoing", "Venture operates independently; xCelero remains a significant shareholder", "Series A+ raised, independent board", "—"],
    ],
    col_widths=[2, 2, 4, 3.5, 2]
)

add_heading_styled("Venture Builder Role", level=2, color=ORANGE)
add_para("Every venture is assigned a Venture Builder — a senior xCelero operator who acts as the interim CEO/COO during stages 1-3, recruits the permanent founding team, and transitions to a board role at stage 4. The Venture Builder is the difference between a venture that ships in 6 months and one that dies in 12.")

add_heading_styled("Domain Sequencing", level=2, color=ORANGE)
add_para("Ventures are sequenced by domain, not built all at once. The sequencing follows three principles: (1) start where the bottleneck is most acute, (2) start where the solution is most deployable, (3) start where the venture can reach revenue fastest. The sequence:")
add_table(
    ["Wave", "Year", "Domains", "Ventures", "Rationale"],
    [
        ["Wave 1", "1–2", "Energy, Water, Finance, Food, Knowledge", "8", "Most acute bottleneck, fastest to revenue"],
        ["Wave 2", "2–3", "Health, Education, Intelligence, Manufacturing, Connectivity", "7", "High impact, proven solution sets"],
        ["Wave 3", "3–4", "Transport, Cities, Science, Materials, Longevity", "7", "Infrastructure-heavy, longer payback"],
        ["Wave 4", "5–7", "Humanoid Robots, Defense, Governance", "5", "Frontier tech, capital-intensive"],
        ["Wave 5", "7–10", "Moon, Mars", "3", "Off-world, decade-long horizons"],
    ],
    col_widths=[2, 2, 4, 2, 4]
)

add_heading_styled("Measurement", level=2, color=ORANGE)
add_bullet("Ventures launched per year (target: 5 in Y1, 50 in Y5, 500 in Y10)")
add_bullet("3-year survival rate (target: 75%)")
add_bullet("Time from ideation to first revenue (target: 6 months)")
add_bullet("Time from first revenue to $1M ARR (target: 18 months)")
add_bullet("Unicorn rate (target: 4% of ventures reach $1Bn valuation)")
add_bullet("Aggregate venture revenue (target: $100Bn by Y10)")

add_page_break()

# ── Chapter 22: Capital ──
add_heading_styled("22. Capital — The Funding Stack", level=1, color=DARK)
add_heading_styled("Owner: Head of Capital (VP-level)", level=2, color=ORANGE)
add_para("Mandate: Build and manage the funding stack that finances ventures and infrastructure. Six vehicles, one thesis: back the technology the next century needs, in the markets that need it most.")

add_heading_styled("The Six Vehicles", level=2, color=ORANGE)
add_table(
    ["Vehicle", "Purpose", "Ticket Size", "Target Close", "Year Launched"],
    [
        ["Studio Fund I–IX", "Core venture-building capital", "$2M → $500M", "Annual", "Y1"],
        ["Thematic Funds", "Domain-specific (Energy, Health, Agri, etc.)", "$5M → $50M", "Bi-annual", "Y2"],
        ["Solidarity Fund", "Democratized access ($500 entry)", "$500 → $5K", "Continuous", "Y1"],
        ["Development Finance", "Concessionary capital for market-entry", "$50K → $1M", "Per-project", "Y2"],
        ["Grant Desk", "Non-dilutive capital matching service", "$10K → $500K", "Continuous", "Y1"],
        ["Secondary Liquidity", "LP/employee liquidity vehicle", "Market", "Opportunistic", "Y5"],
    ],
    col_widths=[3.5, 4, 3, 2.5, 2]
)

add_heading_styled("Solidarity Pricing", level=2, color=ORANGE)
add_para("Early LPs (Years 1-3) receive below-market carry (10% vs. the standard 20%) and a preferential hurdle (8% vs. 10%). This lets the market form before it is squeezed. Once the fund track record is established (Year 4+), pricing normalizes. The principle: the people who believed early share more in the upside.")

add_heading_styled("Capital Deployment Rules", level=2, color=ORANGE)
add_bullet("Maximum 40% of any fund deployed in a single venture.")
add_bullet("Minimum 60% of Studio Fund capital reserved for follow-on (stages 3-5).")
add_bullet("No venture receives more than $2M in total xCelero capital before raising an external round.")
add_bullet("Grant capital is matched 1:1 with equity capital (ensures skin in the game).")

add_page_break()

# ── Chapter 23: Community ──
add_heading_styled("23. Community — The XCitizen Network", level=1, color=DARK)
add_heading_styled("Owner: Head of Community (VP-level)", level=2, color=ORANGE)
add_para("Mandate: Build and nurture the XCitizen network — the operators, founders, investors, and mentors that compound across every hub. The community is the connective tissue that turns isolated efforts into collective momentum.")

add_heading_styled("The XCitizen Journey", level=2, color=ORANGE)
add_table(
    ["Stage", "What Happens", "Duration", "Commitment"],
    [
        ["Applicant", "Applies via website or referral; interviewed by community team", "1–2 weeks", "—"],
        ["Inductee", "Onboarded into a cohort; attends the 8-week XCitizen program", "8 weeks", "10 hrs/week"],
        ["XCitizen", "Full member; access to all hubs, TownSquare, network, programs", "Ongoing", "5 hrs/month minimum"],
        ["Mentor", "XCitizen with 5+ years experience; mentors ventures and fellows", "Ongoing", "4 hrs/month"],
        ["Fellow", "Senior XCitizen; co-leads a hub or a domain vertical", "Ongoing", "10 hrs/month"],
    ],
    col_widths=[2.5, 5, 2.5, 2.5]
)

add_heading_styled("TownSquare — The Digital Forum", level=2, color=ORANGE)
add_para("TownSquare is the digital home of the XCitizen network. It is where the community convenes between cohort residencies, where ventures find mentors and pilots, and where the cross-hub collaborations happen. Target: 1,000,000+ users by Year 10.")
add_bullet("Communities: Energy, Water, Food, Finance, Health, Education, and 14 more — one per domain.")
add_bullet("Features: posts, threaded comments, voting, sharing, AI-assisted posting, real-time updates (Year 2+).")
add_bullet("Governance: community-led moderation, XCitizen-elected community council (Year 3+).")

add_heading_styled("Cohort Cadence", level=2, color=ORANGE)
add_table(
    ["Year", "Cohorts/Year", "XCitizens/Cohort", "Total XCitizens (cumulative)"],
    [
        ["1", "2", "50", "200"],
        ["2", "4", "75", "1,000"],
        ["3", "6", "100", "3,000"],
        ["5", "12", "150", "10,000"],
        ["10", "48", "200", "100,000"],
    ],
    col_widths=[2, 3, 3, 4]
)

add_page_break()

# ── Chapter 24: Programs ──
add_heading_styled("24. Programs — The Four Pipelines", level=1, color=DARK)
add_heading_styled("Owner: Head of Programs (VP-level)", level=2, color=ORANGE)
add_para("Mandate: Operate the four programs that pipeline talent and ventures into the engine. Each program serves a different stage of the founder journey, from raw talent to scaling CEO.")

add_table(
    ["Program", "Stage", "Duration", "Cohort Size", "Output", "Year Launched"],
    [
        ["xHansa Fellowship", "Raw talent → Builder", "8 weeks", "25–50", "Fellows join ventures or XCitizen network", "Y1"],
        ["xCelero Accelerator", "Builder → Founder", "12 weeks", "10–20 ventures", "Validated ventures with first revenue", "Y1"],
        ["Inception Studios", "Idea → ProtoCo", "12 weeks", "5–10 ProtoCos", "ProtoCos ready for NewCo stage", "Y3"],
        ["Quest Fellowship", "Student → Founder", "6 months", "20–40 students", "University ventures spun out", "Y4"],
    ],
    col_widths=[3, 3, 2, 2.5, 3.5, 2]
)

add_heading_styled("Program Economics", level=2, color=ORANGE)
add_bullet("xHansa Fellowship: free for fellows; funded by Studio Fund (community investment).")
add_bullet("xCelero Accelerator: xCelero takes 6% equity; $50K per venture.")
add_bullet("Inception Studios: xCelero takes 10% equity at ProtoCo stage; $25K per ProtoCo.")
add_bullet("Quest Fellowship: co-funded by university partners; xCelero takes 5% equity on spinout.")

add_page_break()

# ── Chapter 25: Technology ──
add_heading_styled("25. Technology — The Platform", level=1, color=DARK)
add_heading_styled("Owner: CTO (VP-level)", level=2, color=ORANGE)
add_para("Mandate: Build and operate the technology platform — the website, TownSquare forum, internal CRM, venture management system, and the data infrastructure that measures the engine.")

add_heading_styled("Platform Roadmap", level=2, color=ORANGE)
add_table(
    ["Year", "Key Platform Milestones"],
    [
        ["1", "Public website, TownSquare v1 (posts, comments, voting), internal CRM, venture tracker"],
        ["2", "TownSquare v2 (search, sharing, AI-assist), mobile app, LP portal, grant-matching system"],
        ["3", "Real-time updates (WebSocket), notifications, moderation tools, analytics dashboard"],
        ["4", "Venture management OS (milestones, KPIs, cap tables), hub utilization tracker"],
        ["5", "AI venture-matching (mentors → ventures), predictive analytics, community governance tools"],
        ["6–10", "Open API (external studios/funds plug in), AI-driven venture ideation, full civilization dashboard"],
    ],
    col_widths=[2, 12]
)

add_heading_styled("Technology Stack", level=2, color=ORANGE)
add_bullet("Frontend: Next.js 16, React, Tailwind CSS, shadcn/ui (already in place)")
add_bullet("Backend: Next.js API routes, Prisma ORM, PostgreSQL (migrate from SQLite in Year 2)")
add_bullet("Real-time: WebSocket mini-service (socket.io, Year 3)")
add_bullet("AI: z-ai-web-dev-sdk (LLM, VLM, TTS, ASR — for AI-assisted posting, venture matching)")
add_bullet("Infrastructure: Vercel (frontend), Railway/Fly.io (backend + DB), Cloudflare (CDN)")
add_bullet("Data: PostgreSQL (operational), ClickHouse (analytics, Year 3+), Redis (cache, Year 2+)")

add_page_break()

# ── Chapter 26: Operations & Talent ──
add_heading_styled("26. Operations & Talent", level=1, color=DARK)
add_heading_styled("Owner: Head of Operations (VP-level)", level=2, color=ORANGE)
add_para("Mandate: Run the operating systems — hiring, finance, legal, admin, and the rhythms (weekly, monthly, quarterly) that keep the engine aligned.")

add_heading_styled("Team Growth", level=2, color=ORANGE)
add_table(
    ["Year", "Headcount (xCelero core)", "Headcount (ventures)", "Total", "Key Hires"],
    [
        ["1", "10", "15", "25", "Founding 6 + 4 ops"],
        ["2", "25", "60", "85", "VP-level across pillars"],
        ["3", "50", "200", "250", "Regional directors"],
        ["5", "150", "2,000", "2,150", "Global functional leaders"],
        ["10", "500", "20,000", "20,500", "Trust governance team"],
    ],
    col_widths=[2, 3, 3, 2, 5]
)

add_heading_styled("Operating Rhythms", level=2, color=ORANGE)
add_bullet("Daily: 15-minute stand-up per team (async on Slack for remote hubs).")
add_bullet("Weekly: All-hands Monday (60 min); venture review Wednesday (90 min).")
add_bullet("Monthly: Board update memo (2 pages, written by CEO); hub managers' call.")
add_bullet("Quarterly: Board meeting; OKR setting; retrospective; plan revision.")
add_bullet("Annually: 10-year plan revision; full team offsite; Demo Day.")

add_page_break()

# ── Chapter 27: Legal, Governance & Compliance ──
add_heading_styled("27. Legal, Governance & Compliance", level=1, color=DARK)
add_heading_styled("Owner: General Counsel (from Year 2; outside counsel in Year 1)", level=2, color=ORANGE)
add_para("Mandate: Ensure xCelero operates legally and ethically across 39 countries, protect the IP and the brand, and govern the transition from studio to trust.")

add_heading_styled("Corporate Structure", level=2, color=ORANGE)
add_bullet("HoldCo: Mauritius (tax-efficient, reputable jurisdiction for African operations).")
add_bullet("OpCos: One per country of operation (Kenya, Nigeria, Ghana, etc.).")
add_bullet("Fund vehicles: Mauritius or Delaware (depending on LP domicile).")
add_bullet("xCelero Trust: established Year 10 as the permanent steward of the engine.")

add_heading_styled("IP Policy", level=2, color=ORANGE)
add_bullet("Venture IP: owned by the venture (xCelero holds equity, not IP).")
add_bullet("Studio IP (shared infrastructure, platform code): owned by xCelero HoldCo.")
add_bullet("Moon/Mars R&D IP: owned by xCelero, licensed to ventures on commercial terms.")

add_heading_styled("Compliance", level=2, color=ORANGE)
add_bullet("Anti-money laundering (AML) and know-your-customer (KYC) for all LPs and venture counterparties.")
add_bullet("Data protection: GDPR-equivalent standards in all markets, regardless of local requirements.")
add_bullet("Anti-corruption: zero-tolerance policy; whistleblower channel; annual training.")

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PART V: METRICS, BUDGETS & RISKS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("PART V: METRICS, BUDGETS & RISKS", level=1, color=ORANGE, size=24)
add_page_break()

# ── Chapter 28: KPI Dashboard ──
add_heading_styled("28. KPI Dashboard — Year by Year", level=1, color=DARK)
add_para("The single source of truth for measuring the engine. Reviewed monthly by the leadership team, quarterly by the board.")

add_heading_styled("Pillar KPIs", level=2, color=ORANGE)
add_table(
    ["KPI", "Y1", "Y3", "Y5", "Y7", "Y10"],
    [
        ["Hubs (cumulative)", "1", "7", "20", "50", "190"],
        ["Countries", "1", "5", "15", "25", "39"],
        ["Ventures (cumulative)", "3", "15", "40", "85", "5,000+"],
        ["Unicorns (cumulative)", "0", "0", "5", "15", "200"],
        ["IPOs (cumulative)", "0", "0", "1", "5", "50"],
        ["Capital raised (cumulative)", "$2M", "$62M", "$322M", "$772M", "$1.97Bn"],
        ["Capital under management", "$2M", "$62M", "$322M", "$772M", "$1.97Bn"],
        ["XCitizens (cumulative)", "200", "3,000", "10,000", "35,000", "100,000"],
        ["TownSquare users", "500", "10,000", "75,000", "300,000", "1,000,000+"],
        ["Jobs created (cumulative)", "15", "200", "2,000", "8,000", "20,000+"],
        ["Domains active", "3", "10", "15", "18", "20"],
        ["Program cohorts (cumulative)", "2", "12", "26", "52", "150+"],
        ["Aggregate venture MRR", "$0", "$500K", "$5M", "$40M", "$8Bn+"],
        ["Aggregate venture revenue (ARR)", "$0", "$6M", "$60M", "$480M", "$100Bn"],
    ],
    col_widths=[5, 2, 2, 2, 2, 2]
)

add_page_break()

# ── Chapter 29: Financial Projections ──
add_heading_styled("29. Financial Projections", level=1, color=DARK)

add_heading_styled("xCelero Labs P&L (Core Operations, USD)", level=2, color=ORANGE)
add_table(
    ["Line Item", "Y1", "Y3", "Y5", "Y7", "Y10"],
    [
        ["Revenue (management fees + carry)", "$200K", "$3M", "$15M", "$40M", "$120M"],
        ["Personnel", "$420K", "$3M", "$12M", "$24M", "$60M"],
        ["Infrastructure (hub operations)", "$200K", "$2M", "$8M", "$16M", "$40M"],
        ["Programs", "$200K", "$2M", "$8M", "$16M", "$40M"],
        ["Technology", "$150K", "$1.5M", "$6M", "$12M", "$30M"],
        ["Other (legal, travel, events)", "$130K", "$1M", "$4M", "$8M", "$20M"],
        ["Total Expenses", "$1.1M", "$9.5M", "$38M", "$76M", "$190M"],
        ["Net Operating Income", "($900K)", "($6.5M)", "($23M)", "($36M)", "($70M)"],
        ["Venture returns (carry, realized)", "$0", "$0", "$2M", "$20M", "$200M"],
        ["Net Income", "($900K)", "($6.5M)", "($21M)", "($16M)", "$130M"],
    ],
    col_widths=[5, 2, 2, 2, 2, 2]
)
add_para("Note: xCelero is designed to operate at a loss for the first 7-8 years, funded by management fees and venture returns. The break-even point is Year 8-9, when accumulated carry from unicorn exits exceeds operating costs. The Decade Fund (Year 10) provides permanent capital that makes the engine self-sustaining.", italic=True, color=GRAY, size=10)

add_heading_styled("Cumulative Capital Requirements", level=2, color=ORANGE)
add_table(
    ["Year", "Operating Capital Needed", "Venture Capital Needed", "Total", "Source"],
    [
        ["1", "$2M", "$0", "$2M", "Studio Fund I"],
        ["2", "$10M", "$5M", "$15M", "Studio Fund II + Thematic"],
        ["3", "$25M", "$20M", "$45M", "Studio Fund III + Thematic"],
        ["5", "$100M", "$100M", "$200M", "Studio Fund V + Thematic"],
        ["7", "$200M", "$300M", "$500M", "Studio Fund VII + Global"],
        ["10", "$500M", "$1Bn", "$1.5Bn", "Decade Fund + all vehicles"],
    ],
    col_widths=[2, 4, 4, 3, 4]
)

add_page_break()

# ── Chapter 30: Milestone Map ──
add_heading_styled("30. The Milestone Map", level=1, color=DARK)
add_para("The non-negotiable milestones — if these are missed, the plan is off-track and the board intervenes.")

add_table(
    ["Milestone", "Target Date", "Consequence if Missed"],
    [
        ["Incorporation + lease signed", "Jan 2026", "Delay everything by 1 quarter"],
        ["XEmbassy Nairobi open", "Apr 2026", "Delay first cohort + first venture"],
        ["First venture (Helios) spun out", "Jun 2026", "Re-evaluate studio model"],
        ["Studio Fund I closed ($2M)", "Jun 2026", "Reduce scope of Year 1"],
        ["First venture revenue (any)", "Dec 2026", "Re-evaluate venture thesis"],
        ["Second hub open (Lagos)", "Q1 2027", "Delay Gulf of Guinea route"],
        ["First $1M ARR venture", "Q2 2028", "Re-evaluate venture scaling"],
        ["First unicorn", "Q3 2029", "Re-evaluate capital model"],
        ["First IPO", "Q1 2030", "Re-evaluate exit strategy"],
        ["100 hubs open", "Q3 2033", "Re-evaluate infrastructure pace"],
        ["200 unicorns", "Q3 2035", "Re-evaluate venture quality"],
        ["190 hubs (complete network)", "Q3 2035", "Extend timeline by 1-2 years"],
        ["xCelero Trust established", "Q4 2035", "Re-evaluate governance model"],
    ],
    col_widths=[6, 3, 6]
)

add_page_break()

# ── Chapter 31: Risk Register ──
add_heading_styled("31. Risk Register", level=1, color=DARK)
add_para("The top risks to the 10-year plan, ranked by likelihood × impact, with mitigation strategies.")

add_table(
    ["#", "Risk", "Likelihood", "Impact", "Mitigation"],
    [
        ["1", "Capital raise fails to meet targets", "Medium", "High", "Diversify LP base; solidarity pricing; milestone-gated tranches"],
        ["2", "Venture survival rate below 75%", "Medium", "High", "Venture Builder model; stage-gated capital; rigorous selection"],
        ["3", "Hub utilization below 80%", "Medium", "Medium", "Hub selection criteria; community-first approach; flexible typology"],
        ["4", "Regulatory barriers in new markets", "High", "Medium", "Local legal counsel; phased entry; government partnerships"],
        ["5", "Key person risk (founder departure)", "Low", "High", "VP-level depth from Year 2; documented processes; board governance"],
        ["6", "Competition from other studios/funds", "Medium", "Medium", "Infrastructure moat; community lock-in; domain specialization"],
        ["7", "Macroeconomic downturn", "Medium", "High", "Counter-cyclical thesis (emerging markets); reserve capital; solidarity pricing"],
        ["8", "Technology platform failure/security breach", "Low", "High", "Postgres migration; Redis caching; security audits; rate limiting"],
        ["9", "Reputational risk (venture scandal/failure)", "Medium", "Medium", "Due diligence; governance; transparency; crisis comms plan"],
        ["10", "Geopolitical instability in operating markets", "High", "Medium", "Route diversification; insurance; exit plans per hub"],
        ["11", "Inability to scale team fast enough", "High", "Medium", "XCitizen pipeline as talent source; remote-first roles; competitive comp"],
        ["12", "Moon/Mars ventures fail to materialize", "High", "Low", "Ring-fenced R&D budget; partnership model; long horizon expectation"],
    ],
    col_widths=[1, 4, 2, 2, 6]
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PART VI: APPENDICES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading_styled("PART VI: APPENDICES", level=1, color=ORANGE, size=24)
add_page_break()

# ── Appendix A: Org Chart ──
add_heading_styled("Appendix A: Organization Chart (Year 3 Target)", level=1, color=DARK)
add_para("The org structure that supports 7 hubs and 15 ventures. The CEO leads; each VP owns a pillar; regional directors manage multi-hub territories.")

add_para("CEO", bold=True, color=ORANGE, size=13)
add_para("  ├── Head of Ventures (VP)", indent=0.5)
add_para("  │     ├── Venture Builders (5, one per active venture)", indent=0.5)
add_para("  │     └── Venture Analyst (1)", indent=0.5)
add_para("  ├── Head of Capital (VP)", indent=0.5)
add_para("  │     ├── Investor Relations (1)", indent=0.5)
add_para("  │     └── Grant Desk Lead (1)", indent=0.5)
add_para("  ├── Head of Infrastructure (VP)", indent=0.5)
add_para("  │     ├── Regional Director, East Africa (1)", indent=0.5)
add_para("  │     ├── Regional Director, West Africa (1)", indent=0.5)
add_para("  │     └── Hub Managers (7, one per hub)", indent=0.5)
add_para("  ├── Head of Community (VP)", indent=0.5)
add_para("  │     ├── Community Managers (3)", indent=0.5)
add_para("  │     └── Program Managers (4, one per program)", indent=0.5)
add_para("  ├── CTO", indent=0.5)
add_para("  │     ├── Frontend Engineer (2)", indent=0.5)
add_para("  │     ├── Backend Engineer (2)", indent=0.5)
add_para("  │     └── Designer (1)", indent=0.5)
add_para("  ├── Head of Operations (VP / COO)", indent=0.5)
add_para("  │     ├── Finance Manager (1)", indent=0.5)
add_para("  │     ├── HR / People (1)", indent=0.5)
add_para("  │     └── Admin (2)", indent=0.5)
add_para("  └── General Counsel (1)", indent=0.5)
add_para("Total: ~40 core staff (Year 3)", italic=True, color=GRAY)

add_page_break()

# ── Appendix B: Hub Rollout Schedule ──
add_heading_styled("Appendix B: Hub Rollout Schedule (190 Hubs)", level=1, color=DARK)
add_para("The complete 190-hub rollout, by route and year. Each hub is listed with its city, type (M1/M2/M3), and target opening year.")

add_heading_styled("Route 1: Gulf of Guinea Arc (Hubs 1–40)", level=2, color=ORANGE)
add_table(
    ["#", "City", "Country", "Type", "Year"],
    [
        ["1", "Nairobi", "Kenya", "M1", "2026"],
        ["2", "Lagos", "Nigeria", "M1", "2027"],
        ["3", "Accra", "Ghana", "M1", "2027"],
        ["4", "Kano", "Nigeria", "M2", "2028"],
        ["5", "Abidjan", "Côte d'Ivoire", "M2", "2028"],
        ["6–10", "Port Harcourt, Kumasi, Cotonou, Lomé, Douala", "Nigeria/Ghana/Benin/Togo/Cameroon", "M2/M3", "2029–2030"],
        ["11–20", "West African secondary cities", "Various", "M2/M3", "2031–2033"],
        ["21–40", "West African tertiary cities", "Various", "M3", "2034–2035"],
    ],
    col_widths=[2, 4, 3, 2, 2]
)

add_heading_styled("Route 2: Sahel Corridor (Hubs 41–70)", level=2, color=ORANGE)
add_table(
    ["#", "City", "Country", "Type", "Year"],
    [
        ["41", "Addis Ababa", "Ethiopia", "M1", "2028"],
        ["42", "Kampala", "Uganda", "M2", "2028"],
        ["43", "Niamey", "Niger", "M2", "2029"],
        ["44", "Bamako", "Mali", "M2", "2029"],
        ["45", "Ouagadougou", "Burkina Faso", "M3", "2030"],
        ["46–70", "Sahel secondary + tertiary cities", "Various", "M2/M3", "2031–2035"],
    ],
    col_widths=[2, 4, 3, 2, 2]
)

add_heading_styled("Route 3: East African Corridor (Hubs 71–100)", level=2, color=ORANGE)
add_table(
    ["#", "City", "Country", "Type", "Year"],
    [
        ["71", "Dar es Salaam", "Tanzania", "M1", "2029"],
        ["72", "Kigali", "Rwanda", "M2", "2029"],
        ["73", "Lusaka", "Zambia", "M2", "2029"],
        ["74–100", "East African secondary + tertiary cities", "Various", "M2/M3", "2030–2035"],
    ],
    col_widths=[2, 4, 3, 2, 2]
)

add_heading_styled("Route 4: Southern Africa (Hubs 101–130)", level=2, color=ORANGE)
add_table(
    ["#", "City", "Country", "Type", "Year"],
    [
        ["101", "Cape Town", "South Africa", "M1", "2029"],
        ["102", "Johannesburg", "South Africa", "M1", "2029"],
        ["103", "Harare", "Zimbabwe", "M2", "2030"],
        ["104", "Maputo", "Mozambique", "M2", "2030"],
        ["105–130", "Southern African secondary + tertiary cities", "Various", "M2/M3", "2031–2035"],
    ],
    col_widths=[2, 4, 3, 2, 2]
)

add_heading_styled("Route 5: North Africa (Hubs 131–160)", level=2, color=ORANGE)
add_table(
    ["#", "City", "Country", "Type", "Year"],
    [
        ["131", "Cairo", "Egypt", "M1", "2030"],
        ["132", "Tunis", "Tunisia", "M1", "2030"],
        ["133", "Algiers", "Algeria", "M2", "2030"],
        ["134", "Casablanca", "Morocco", "M1", "2030"],
        ["135–160", "North African secondary + tertiary cities", "Various", "M2/M3", "2031–2035"],
    ],
    col_widths=[2, 4, 3, 2, 2]
)

add_heading_styled("Route 6: Global (Hubs 161–190)", level=2, color=ORANGE)
add_table(
    ["#", "City", "Country/Region", "Type", "Year"],
    [
        ["161", "Dubai", "UAE", "M1", "2031"],
        ["162", "Mumbai", "India", "M1", "2031"],
        ["163", "Singapore", "Singapore", "M1", "2031"],
        ["164", "São Paulo", "Brazil", "M1", "2032"],
        ["165", "Mexico City", "Mexico", "M1", "2032"],
        ["166–190", "Global secondary cities (Asia, LatAm, ME)", "Various", "M2/M3", "2033–2035"],
    ],
    col_widths=[2, 4, 3, 2, 2]
)

add_page_break()

# ── Appendix C: Fund Structure ──
add_heading_styled("Appendix C: Fund Structure & Capital Stack", level=1, color=DARK)
add_para("The full capital stack across all six vehicles, showing how capital flows from LPs through xCelero to ventures and infrastructure.")

add_heading_styled("Capital Stack (Year 10 Target)", level=2, color=ORANGE)
add_table(
    ["Vehicle", "Total Capital (Y10)", "LPs", "Carry", "Hurdle", "Deployed Into"],
    [
        ["Studio Funds I–IX", "$1.2Bn", "500+", "20% (10% early)", "8%", "Ventures (stages 1–5)"],
        ["Thematic Funds (8)", "$400M", "200+", "20%", "10%", "Domain-specific ventures"],
        ["Solidarity Fund", "$50M", "5,000+", "10%", "5%", "Early-stage ventures"],
        ["Development Finance", "$150M", "DFIs/govts", "0% (concessionary)", "0%", "Market-entry pilots"],
        ["Grant Desk (matched)", "$100M", "Foundations/govts", "0%", "0%", "Non-dilutive venture grants"],
        ["Secondary Liquidity", "$75M", "Institutional", "10%", "0%", "LP/employee buyouts"],
        ["Decade Fund (permanent)", "$500M", "Anchor LPs + Trust", "5%", "4%", "Permanent capital"],
        ["Total", "$2.475Bn", "—", "—", "—", "—"],
    ],
    col_widths=[4, 3, 2, 2, 2, 4]
)

add_page_break()

# ── Appendix D: The Twenty Domains Deep Dive ──
add_heading_styled("Appendix D: The Twenty Domains — Deep Dive", level=1, color=DARK)
add_para("For each of the twenty domains: the bottleneck, the solution set, the first venture xCelero builds, the KPI, and the target year of deployment.")

domains = [
    ("Food", "More food needed in 40 years than in 8,000", "Vertical farms, AI crops, lab protein", "Bounty (XCL-202)", "10× yield per drop", "Y1"),
    ("Energy", "1.3B without power; diesel skies", "Fusion, solar paint, peer-to-peer grids", "Helios (XCL-101)", "10,000+ microgrid nodes", "Y1"),
    ("Water", "2B drink sick water; child dies every 80s", "Atmospheric generators, low-energy desal", "Nimbus (XCL-105)", "$0.01/L clean water", "Y1"),
    ("Finance", "1.4B unbanked; 15% remittance fees", "Money as internet protocol; instant", "Meridian (XCL-203)", "Instant, borderless, <$0.01/txn", "Y1"),
    ("Knowledge", "Talent sealed off by birthplace", "Single network; talent finds capital in a day", "XCitizen Network", "100,000 XCitizens by Y10", "Y1"),
    ("Health", "1 in 5 have diagnostics", "Hospital of the future: sensor + AI", "Refract (XCL-301)", "Diagnostics in every ward", "Y2"),
    ("Education", "617M illiterate", "1:1 AI tutor per child, free, $20 phone", "Ceres Learn (XCL-401)", "1:1 tutor for 100M children", "Y2"),
    ("Intelligence", "Frontier model costs $100M", "Compute as utility; AI on every phone", "EdgeMind (XCL-501)", "Frontier AI on $20 phone", "Y2"),
    ("Manufacturing", "Sell lithium, buy back battery at 10×", "Micro-factories; 12-block supply chain", "Manna (XCL-204)", "Local micro-factory in every town", "Y2"),
    ("Connectivity", "2.6B offline", "Satellite blanket; last billion online", "Constella (XCL-601)", "2.6B online", "Y2"),
    ("Transport", "75% of product cost is logistics", "eVTOL, hyperloop, one multimodal OS", "Velocity (XCL-701)", "Point-to-point in minutes", "Y3"),
    ("Cities", "70% urban by 2050; slums fastest-growing", "Cities that grow like coral", "Civitas (XCL-801)", "Self-powering, self-cleaning cities", "Y3"),
    ("Science", "0.1% of GDP on basic research", "AI runs 1B experiments in simulation", "LabOS (XCL-901)", "Fusion breakthrough in 6 months", "Y3"),
    ("Longevity", "73 yrs avg; $1T/yr managing decline", "Cellular reprogramming, senolytics", "Allele (XCL-1001)", "120+ as baseline", "Y4"),
    ("Materials", "Mine-to-landfill one-way pipe", "Atom-by-atom design; mine the landfill", "Forge (XCL-1101)", "Circular material economy", "Y4"),
    ("Humanoid Robots", "$100T/yr of manual human effort", "$20K robot per home; work optional", "Titan (XCL-1201)", "1 robot per home", "Y5"),
    ("Defense", "$2.4T/yr buying hardware for last war", "Defense as software; war irrational", "Aegis (XCL-1301)", "Cost of aggression irrational", "Y5"),
    ("Governance", "17th-century state for satellite planet", "Governance as protocol; opt-in communities", "CivicID v2 (XCL-1401)", "Fluid, opt-in governance", "Y6"),
    ("The Moon", "12 walked it; none since 1972", "Permanent settlement; industry off-world", "Selene (XCL-1501)", "Permanent lunar settlement", "Y7"),
    ("Mars", "0 humans; 30M-mile desert", "Self-sustaining city of 1M", "Ares (XCL-1601)", "City of 1M on Mars", "Y8"),
]

add_table(
    ["Domain", "Bottleneck", "Solution", "First Venture", "KPI", "Year"],
    [list(d) for d in domains],
    col_widths=[2.5, 3, 3, 2.5, 3, 1]
)

add_page_break()

# ── Appendix E: Glossary ──
add_heading_styled("Appendix E: Glossary of Terms", level=1, color=DARK)

glossary = [
    ("XEmbassy", "A physical xCelero hub — a campus containing some combination of foundry, lab, studio, commons, and outpost."),
    ("XCitizen", "A member of the xCelero community network — an operator, founder, investor, or mentor who has completed the onboarding program."),
    ("ProtoCo", "A venture in the prototype stage (stage 2 of the venture pipeline) — has an MVP but no revenue."),
    ("NewCo", "A venture that has incorporated and is deploying its pilot (stage 3)."),
    ("Studio Fund", "The core venture-building capital vehicle. Annual vintage, funds venture creation across all domains."),
    ("Thematic Fund", "A domain-specific capital vehicle (e.g., Energy Thematic Fund) that co-invests alongside the Studio Fund."),
    ("Solidarity Pricing", "Below-market carry and hurdle for early LPs, designed to let the market form before it is squeezed."),
    ("Flywheel", "The compounding loop: infrastructure → ventures → capital → community → infrastructure. Each turn is faster than the last."),
    ("Route", "A geographic corridor along which xCelero builds hubs (e.g., Gulf of Guinea Arc, Sahel Corridor)."),
    ("Venture Builder", "A senior xCelero operator who acts as interim CEO/COO during a venture's early stages, then transitions to a board role."),
    ("TownSquare", "The digital forum where the XCitizen network convenes. Posts, comments, voting, sharing, AI-assisted posting."),
    ("Decade Fund", "The permanent capital vehicle established in Year 10 to sustain the engine in perpetuity."),
    ("xCelero Trust", "The governance entity established in Year 10 to steward the engine as a permanent institution."),
    ("Domain", "One of the twenty civilization pillars xCelero builds ventures in (Food, Energy, Water, etc.)."),
    ("Strike Zone", "A geographic area where an XEmbassy is placed and active (synonym for a hub location)."),
    ("Demo Day", "The culminating event of a program cohort, where ventures present to investors and the community."),
    ("Inception Studios", "The xCelero program that takes ideas from inception to ProtoCo stage."),
    ("xHansa Fellowship", "The xCelero program that develops raw talent into builders (8-week crucible)."),
    ("xCelero Accelerator", "The xCelero program that takes builders to founders (12-week, 6% equity, $50K)."),
    ("Quest Fellowship", "The xCelero university partnership program (6-month, student founders)."),
]

add_table(
    ["Term", "Definition"],
    [[t, d] for t, d in glossary],
    col_widths=[4, 12]
)

# ── Closing ──
add_page_break()
add_divider()
add_para("END OF DOCUMENT", bold=True, color=ORANGE, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("xCelero Labs — 10-Year Workplan (2026–2035)", italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Version 1.0 — June 2026", color=LIGHT_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_para("This is a living document. It will be revised quarterly. The Year 1 detail is prescriptive; the Year 10 detail is directional. The further out we go, the more the plan becomes a compass rather than a map.", italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_divider()

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
output_path = "/home/z/my-project/download/xCelero_Labs_10_Year_Workplan.docx"
doc.save(output_path)
print(f"✅ Document saved: {output_path}")
print(f"   File size: {os.path.getsize(output_path) / 1024:.0f} KB")

# Count paragraphs and tables for a rough page estimate
para_count = len(doc.paragraphs)
table_count = len(doc.tables)
print(f"   Paragraphs: {para_count}")
print(f"   Tables: {table_count}")
print(f"   Estimated pages: ~{para_count // 8 + table_count * 0.5:.0f}")
